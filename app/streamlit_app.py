import base64
import json
import html
import os
import re
from io import BytesIO
from urllib import error, request

import streamlit as st
import pandas as pd
import requests
from pathlib import Path
from datetime import datetime

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_DIR / "outputs"
DATA_DIR = PROJECT_DIR / "data"
POLICY_DIR = PROJECT_DIR / "policy"


def get_config_value(name, default=""):
    value = os.getenv(name)
    if value:
        return value
    try:
        secret_value = st.secrets.get(name, default)
    except Exception:
        secret_value = default
    return str(secret_value).strip() if secret_value is not None else default

DHM_IMG = OUTPUT_DIR / "dhm_result_clean.png"
HLG_IMG = OUTPUT_DIR / "hlg_result_clean.png"
DHM_SUMMARY = DATA_DIR / "dhm_summary.csv"
HLG_SUMMARY = DATA_DIR / "hlg_summary.csv"
DHM_GEOJSON = DATA_DIR / "dhm.geojson"
HLG_GEOJSON = DATA_DIR / "hlg.geojson"

SCENARIOS = ["候鸟老人友好模式", "年轻家庭模式", "游客短租模式"]

LAND_USE_LEGEND = [
    {"编号": 1, "上游类型": "feasible", "展示端颜色": "蓝色", "色值": "#1f77b4", "类型名称": "可规划地块 / 剩余可建设空间", "说明": "GeoJSON type=1，对应可规划地块或剩余可建设空间。"},
    {"编号": 2, "上游类型": "road", "展示端颜色": "蓝色", "色值": "#1f77b4", "类型名称": "道路", "说明": "GeoJSON type=2，对应道路。"},
    {"编号": 3, "上游类型": "boundary", "展示端颜色": "橙色", "色值": "#ff7f0e", "类型名称": "边界", "说明": "GeoJSON type=3，对应边界。"},
    {"编号": 4, "上游类型": "residential", "展示端颜色": "绿色", "色值": "#2ca02c", "类型名称": "居住用地", "说明": "GeoJSON type=4，对应居住用地。"},
    {"编号": 5, "上游类型": "business", "展示端颜色": "红色", "色值": "#d62728", "类型名称": "商业用地", "说明": "GeoJSON type=5，对应商业用地。"},
    {"编号": 6, "上游类型": "office", "展示端颜色": "紫色", "色值": "#9467bd", "类型名称": "办公用地", "说明": "GeoJSON type=6，对应办公用地。"},
    {"编号": 7, "上游类型": "green_l", "展示端颜色": "棕色", "色值": "#8c564b", "类型名称": "大型绿地", "说明": "GeoJSON type=7，对应大型绿地。"},
    {"编号": 8, "上游类型": "green_s", "展示端颜色": "棕色", "色值": "#8c564b", "类型名称": "小型绿地 / 口袋绿地", "说明": "GeoJSON type=8，对应小型绿地或口袋绿地。"},
    {"编号": 9, "上游类型": "school", "展示端颜色": "粉色", "色值": "#e377c2", "类型名称": "学校 / 教育设施", "说明": "GeoJSON type=9，对应学校或教育设施。"},
    {"编号": 10, "上游类型": "hospital_l", "展示端颜色": "灰色", "色值": "#7f7f7f", "类型名称": "大型医院 / 综合医疗", "说明": "GeoJSON type=10，对应大型医院或综合医疗。"},
    {"编号": 11, "上游类型": "hospital_s", "展示端颜色": "黄绿色", "色值": "#bcbd22", "类型名称": "小型医院 / 社区医疗", "说明": "GeoJSON type=11，对应小型医院或社区医疗。"},
    {"编号": 12, "上游类型": "recreation", "展示端颜色": "青色", "色值": "#17becf", "类型名称": "休闲娱乐 / 文体活动", "说明": "GeoJSON type=12，对应休闲娱乐或文体活动。"},
    {"编号": 13, "上游类型": "intersection", "展示端颜色": "青色圆点", "色值": "#17becf", "类型名称": "道路交叉点 / 路网节点", "说明": "GeoJSON type=13，对应道路交叉点或路网节点。"},
]

LAND_USE_COLOR_MAP = {item["编号"]: item["色值"] for item in LAND_USE_LEGEND}

SCENE_FOCUS_TYPES = {
    "候鸟老人友好模式": {10, 11, 7, 8, 13},
    "年轻家庭模式": {9, 4, 7, 8, 2},
    "游客短租模式": {5, 12, 2, 13},
}

INSTANT_PLAN_CONFIG = {
    "候鸟老人友好模式": {
        "focus_types": {10, 11, 7, 8, 13},
        "suggestions": [
            "建议增设休息座椅",
            "优化无障碍慢行路径",
            "优先保障社区医疗可达",
        ],
    },
    "年轻家庭模式": {
        "focus_types": {9, 4, 8, 2},
        "suggestions": [
            "提升学校与居住组团联系",
            "增加亲子活动空间",
            "控制居住区周边交通干扰",
        ],
    },
    "游客短租模式": {
        "focus_types": {5, 12, 2, 13},
        "suggestions": [
            "强化公交接驳与景点联系",
            "引导短租服务集中布局",
            "夜间出行安全节点优化",
        ],
    },
}

RULE_PARSE_DEFAULTS = {
    "候鸟老人友好模式": {
        "medical_weight": 0.35,
        "education_weight": 0.00,
        "commerce_weight": 0.25,
        "green_weight": 0.25,
        "traffic_weight": 0.15,
        "walking_time": 10,
        "explanation": "命中老人、候鸟、养老、医院、买菜或公园等关键词，按候鸟老人友好场景配置权重。",
    },
    "年轻家庭模式": {
        "medical_weight": 0.05,
        "education_weight": 0.35,
        "commerce_weight": 0.25,
        "green_weight": 0.20,
        "traffic_weight": 0.15,
        "walking_time": 15,
        "explanation": "命中孩子、学校、教育、家庭、幼儿园或上学等关键词，按年轻家庭场景配置权重。",
    },
    "游客短租模式": {
        "medical_weight": 0.05,
        "education_weight": 0.05,
        "commerce_weight": 0.30,
        "green_weight": 0.15,
        "traffic_weight": 0.45,
        "walking_time": 15,
        "explanation": "命中游客、短租、免税、景点、酒店、交通或民宿等关键词，按游客短租场景配置权重。",
    },
}

WEIGHT_FIELDS = [
    "medical_weight",
    "education_weight",
    "commerce_weight",
    "green_weight",
    "traffic_weight",
]


def extract_walking_time(user_text, default_time):
    match = re.search(r"(\d{1,2})\s*分钟", user_text)
    if match:
        return int(match.group(1))
    return default_time


def parse_with_rules(user_text, selected_scene):
    text = user_text.lower()
    scenario = selected_scene if selected_scene in SCENARIOS else SCENARIOS[0]

    if any(k in text for k in ["老人", "候鸟", "养老", "慢病", "医院", "买菜", "公园"]):
        scenario = "候鸟老人友好模式"
    elif any(k in text for k in ["孩子", "学校", "教育", "家庭", "幼儿园", "上学"]):
        scenario = "年轻家庭模式"
    elif any(k in text for k in ["游客", "短租", "免税", "景点", "酒店", "交通", "民宿"]):
        scenario = "游客短租模式"

    result = {
        "scenario": scenario,
        **RULE_PARSE_DEFAULTS[scenario],
        "parse_method": "规则解析",
    }
    result["walking_time"] = extract_walking_time(user_text, result["walking_time"])
    return normalize_weights(result, RULE_PARSE_DEFAULTS[scenario])


def coerce_weight(value, fallback):
    if value is None or value == "":
        value = fallback
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = float(fallback)
    if number > 1:
        number = number / 100
    return max(0.0, min(1.0, number))


def normalize_weights(result, fallback):
    normalized = dict(result)
    scenario = normalized.get("scenario")
    default_weights = RULE_PARSE_DEFAULTS.get(scenario, fallback)
    weights = []

    for field in WEIGHT_FIELDS:
        weights.append(coerce_weight(normalized.get(field, default_weights[field]), default_weights[field]))

    total = sum(weights)
    if total <= 0:
        weights = [coerce_weight(default_weights[field], RULE_PARSE_DEFAULTS[SCENARIOS[0]][field]) for field in WEIGHT_FIELDS]
        total = sum(weights)

    if total <= 0:
        default_weights = RULE_PARSE_DEFAULTS[SCENARIOS[0]]
        weights = [default_weights[field] for field in WEIGHT_FIELDS]
        total = sum(weights)

    for field, value in zip(WEIGHT_FIELDS, weights):
        normalized[field] = value / total

    return normalized


def render_land_use_legend_table(focus_types=None, include_focus=False):
    focus_types = set(focus_types or [])
    return pd.DataFrame(
        [
            {
                "type": item["编号"],
                "上游类型": item["上游类型"],
                "中文含义": item["类型名称"],
                "展示端颜色": item["展示端颜色"],
                **({"当前场景重点": "是" if item["编号"] in focus_types else ""} if include_focus else {}),
                "说明": item["说明"],
            }
            for item in LAND_USE_LEGEND
        ]
    )


def render_land_use_legend_html(focus_types=None, include_focus=False):
    focus_types = set(focus_types or [])
    headers = ["type", "上游类型", "中文含义", "展示端颜色"]
    if include_focus:
        headers.append("当前场景重点")
    headers.append("说明")

    header_html = "".join([f"<th>{html.escape(header)}</th>" for header in headers])
    row_html = []

    for item in LAND_USE_LEGEND:
        land_type = item["编号"]
        color_value = item["色值"]
        swatch_class = "legend-swatch legend-dot" if land_type == 13 else "legend-swatch"
        color_html = (
            f'<span class="{swatch_class}" style="background:{html.escape(color_value)};"></span>'
            f"{html.escape(item['展示端颜色'])}"
        )
        cells = [
            f'<td class="land-use-type">{land_type}</td>',
            f"<td>{html.escape(item['上游类型'])}</td>",
            f"<td>{html.escape(item['类型名称'])}</td>",
            f'<td class="legend-color-cell">{color_html}</td>',
        ]
        if include_focus:
            focus_html = '<span class="legend-focus-badge">重点</span>' if land_type in focus_types else '<span class="legend-focus-muted">-</span>'
            cells.append(f'<td class="legend-focus-cell">{focus_html}</td>')
        cells.append(f"<td>{html.escape(item['说明'])}</td>")
        row_html.append(f"<tr>{''.join(cells)}</tr>")

    return f"""
<div class="land-use-table-wrap">
  <table class="land-use-table">
    <thead><tr>{header_html}</tr></thead>
    <tbody>{''.join(row_html)}</tbody>
  </table>
</div>
"""


def format_summary_table_value(value):
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.2f}"
    return str(value)


def render_summary_table_html(dataframe):
    preferred_columns = ["type", "geom_type", "count", "total_area", "total_length"]
    columns = [column for column in preferred_columns if column in dataframe.columns]
    if not columns:
        columns = dataframe.columns.tolist()

    header_html = "".join([f"<th>{html.escape(str(column))}</th>" for column in columns])
    rows = []
    for _, row in dataframe[columns].iterrows():
        cells = [
            f"<td>{html.escape(format_summary_table_value(row[column]))}</td>"
            for column in columns
        ]
        rows.append(f"<tr>{''.join(cells)}</tr>")

    return f"""
<div class="summary-table-wrap">
  <table class="summary-table">
    <thead><tr>{header_html}</tr></thead>
    <tbody>{''.join(rows)}</tbody>
  </table>
</div>
"""


def normalize_geojson_type(value):
    try:
        if value is None:
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def get_geojson_feature_type(feature):
    properties = feature.get("properties", {}) if isinstance(feature, dict) else {}
    for field in ["type", "TYPE", "Type"]:
        if field in properties:
            return normalize_geojson_type(properties.get(field))
    return None


def extract_xy_sequence(coordinates):
    points = []
    for point in coordinates or []:
        if isinstance(point, (list, tuple)) and len(point) >= 2:
            points.append((float(point[0]), float(point[1])))
    return points


def draw_geojson_polygon(ax, rings, color, is_focus):
    exterior = extract_xy_sequence(rings[0] if rings else [])
    if len(exterior) < 3:
        return
    xs, ys = zip(*exterior)
    ax.fill(
        xs,
        ys,
        facecolor=color,
        edgecolor="#17324d" if is_focus else "#2f3f4f",
        linewidth=1.35 if is_focus else 0.35,
        alpha=0.9 if is_focus else 0.48,
        zorder=5 if is_focus else 2,
    )


def draw_geojson_line(ax, coordinates, color, is_focus):
    points = extract_xy_sequence(coordinates)
    if len(points) < 2:
        return
    xs, ys = zip(*points)
    ax.plot(
        xs,
        ys,
        color=color,
        linewidth=2.4 if is_focus else 0.85,
        alpha=0.95 if is_focus else 0.7,
        solid_capstyle="round",
        zorder=8 if is_focus else 6,
    )


def draw_geojson_point(ax, coordinates, color, is_focus):
    point = extract_xy_sequence([coordinates])
    if not point:
        return
    ax.scatter(
        [point[0][0]],
        [point[0][1]],
        s=28 if is_focus else 9,
        color=color,
        edgecolors="#17324d" if is_focus else "none",
        linewidths=0.9 if is_focus else 0,
        alpha=0.95 if is_focus else 0.72,
        zorder=10 if is_focus else 7,
    )


def draw_geojson_geometry(ax, geometry, color, is_focus):
    if not isinstance(geometry, dict):
        return
    geom_type = geometry.get("type")
    coordinates = geometry.get("coordinates")
    if geom_type == "Polygon":
        draw_geojson_polygon(ax, coordinates, color, is_focus)
    elif geom_type == "MultiPolygon":
        for polygon in coordinates or []:
            draw_geojson_polygon(ax, polygon, color, is_focus)
    elif geom_type == "LineString":
        draw_geojson_line(ax, coordinates, color, is_focus)
    elif geom_type == "MultiLineString":
        for line in coordinates or []:
            draw_geojson_line(ax, line, color, is_focus)
    elif geom_type == "Point":
        draw_geojson_point(ax, coordinates, color, is_focus)
    elif geom_type == "MultiPoint":
        for point in coordinates or []:
            draw_geojson_point(ax, point, color, is_focus)


def collect_geometry_points(geometry):
    points = []

    def walk(value):
        if isinstance(value, (list, tuple)):
            if len(value) >= 2 and all(isinstance(value[index], (int, float)) for index in [0, 1]):
                points.append((float(value[0]), float(value[1])))
            else:
                for item in value:
                    walk(item)

    if isinstance(geometry, dict):
        walk(geometry.get("coordinates"))
    return points


def get_feature_center(feature):
    points = collect_geometry_points(feature.get("geometry") if isinstance(feature, dict) else None)
    if not points:
        return None
    xs, ys = zip(*points)
    return (sum(xs) / len(xs), sum(ys) / len(ys))


def get_map_bounds(features):
    points = []
    for feature in features:
        points.extend(collect_geometry_points(feature.get("geometry") if isinstance(feature, dict) else None))
    if not points:
        return None
    xs, ys = zip(*points)
    return min(xs), min(ys), max(xs), max(ys)


def render_geojson_plan_figure(geojson_path, focus_types):
    if not MATPLOTLIB_AVAILABLE:
        raise RuntimeError("matplotlib is not installed")
    if not geojson_path.exists() or geojson_path.stat().st_size == 0:
        raise FileNotFoundError(str(geojson_path))

    with open(geojson_path, "r", encoding="utf-8") as file:
        geojson = json.load(file)

    features = geojson.get("features", [])
    if not isinstance(features, list) or not features:
        raise ValueError("GeoJSON has no features")

    focus_types = set(focus_types or [])
    fig, ax = plt.subplots(figsize=(4.8, 5.2), dpi=140)
    fig.patch.set_facecolor("#fffdf8")
    ax.set_facecolor("#fffdf8")
    sorted_features = sorted(
        features,
        key=lambda feature: get_geojson_feature_type(feature) in focus_types,
    )

    for feature in sorted_features:
        land_use_type = get_geojson_feature_type(feature)
        if land_use_type is None:
            continue
        color = LAND_USE_COLOR_MAP.get(land_use_type, "#9ca3af")
        is_focus = land_use_type in focus_types
        draw_geojson_geometry(ax, feature.get("geometry"), color, is_focus)

    ax.set_aspect("equal", adjustable="box")
    ax.autoscale()
    ax.margins(0.01)
    ax.axis("off")
    fig.tight_layout(pad=0)
    return fig, len(features)


def get_instant_plan_config(scene_name):
    return INSTANT_PLAN_CONFIG.get(scene_name, INSTANT_PLAN_CONFIG["候鸟老人友好模式"])


def render_instant_suggestion_cards(suggestions):
    cards = []
    for index, suggestion in enumerate(suggestions, start=1):
        cards.append(
            f"""
<div class="instant-suggestion-card">
  <div class="instant-suggestion-index">{index}</div>
  <h4>{html.escape(suggestion)}</h4>
  <p>图中编号 {index} 为该建议的示意标注位置。</p>
</div>
"""
        )
    st.markdown(
        f"""
<div class="instant-suggestion-grid">
{''.join(cards)}
</div>
""",
        unsafe_allow_html=True,
    )


def get_focus_type_names(focus_type_ids):
    names = []
    for item in LAND_USE_LEGEND:
        if item["编号"] in set(focus_type_ids):
            names.append(f"type={item['编号']} {item['类型名称']}")
    return names


def render_instant_plan_report(scene_name, parse_result):
    config = get_instant_plan_config(scene_name)
    focus_names = get_focus_type_names(config["focus_types"])
    suggestions_text = "\n".join([f"{index}. {item}" for index, item in enumerate(config["suggestions"], start=1)])
    focus_text = "、".join(focus_names) if focus_names else "当前场景重点类型"
    weights_text = (
        f"医疗{parse_result['medical_weight']:.0%}、教育{parse_result['education_weight']:.0%}、"
        f"商业{parse_result['commerce_weight']:.0%}、绿地{parse_result['green_weight']:.0%}、"
        f"交通{parse_result['traffic_weight']:.0%}"
    )
    return f"""## 即时规划建议说明

- 当前场景：{scene_name}
- 指标权重参考：{weights_text}
- 重点关注类型：{focus_text}

### 三条即时建议

{suggestions_text}

真实性边界说明：该建议为展示端基于离线 GeoJSON 结果和当前场景的即时标注说明，用于解释用户需求下的优化方向；当前版本不进行现场 PPO/SGNN 训练，也不实时生成新的底层规划结果。"""


def render_instant_plan_figure(geojson_path, scene_name, parse_result):
    if not MATPLOTLIB_AVAILABLE:
        raise RuntimeError("matplotlib is not installed")
    if not geojson_path.exists() or geojson_path.stat().st_size == 0:
        raise FileNotFoundError(str(geojson_path))

    with open(geojson_path, "r", encoding="utf-8") as file:
        geojson = json.load(file)

    features = geojson.get("features", [])
    if not isinstance(features, list) or not features:
        raise ValueError("GeoJSON has no features")

    instant_config = get_instant_plan_config(scene_name)
    focus_types = set(instant_config["focus_types"])
    fig, ax = plt.subplots(figsize=(5.2, 5.4), dpi=150)
    fig.patch.set_facecolor("#fffdf8")
    ax.set_facecolor("#fffdf8")

    sorted_features = sorted(
        features,
        key=lambda feature: get_geojson_feature_type(feature) in focus_types,
    )
    for feature in sorted_features:
        land_use_type = get_geojson_feature_type(feature)
        if land_use_type is None:
            continue
        color = LAND_USE_COLOR_MAP.get(land_use_type, "#9ca3af")
        is_focus = land_use_type in focus_types
        draw_geojson_geometry(ax, feature.get("geometry"), color, is_focus)

    bounds = get_map_bounds(features)
    if bounds:
        min_x, min_y, max_x, max_y = bounds
        width = max(max_x - min_x, 1e-6)
        height = max(max_y - min_y, 1e-6)
        marker_positions = [
            (min_x + width * 0.25, min_y + height * 0.76),
            (min_x + width * 0.68, min_y + height * 0.54),
            (min_x + width * 0.43, min_y + height * 0.24),
        ]
    else:
        marker_positions = [(0, 0), (1, 1), (2, 0)]

    marker_colors = ["#1b8fb8", "#18a99d", "#ff8f7a"]
    for index, (x_value, y_value) in enumerate(marker_positions, start=1):
        ax.scatter(
            [x_value],
            [y_value],
            s=220,
            color=marker_colors[index - 1],
            edgecolors="#ffffff",
            linewidths=1.8,
            alpha=0.94,
            zorder=30,
        )
        ax.text(
            x_value,
            y_value,
            str(index),
            color="#ffffff",
            ha="center",
            va="center",
            fontsize=11,
            fontweight="bold",
            zorder=31,
        )

    ax.set_aspect("equal", adjustable="box")
    ax.autoscale()
    ax.margins(0.025)
    ax.axis("off")
    fig.tight_layout(pad=0.05)
    return fig, instant_config


def render_matplotlib_figure_png(fig):
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=140, bbox_inches="tight", pad_inches=0.02, facecolor=fig.get_facecolor())
    buffer.seek(0)
    return buffer.getvalue()


def render_centered_image(image_or_bytes, caption=None, width=520):
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image(image_or_bytes, caption=caption, width=width)


def render_planning_image_scroll(image_path):
    image_base64 = base64.b64encode(image_path.read_bytes()).decode("ascii")
    st.markdown(
        f"""
<div class="planning-image-scroll">
  <img src="data:image/png;base64,{image_base64}" alt="规划结果图" />
</div>
""",
        unsafe_allow_html=True,
    )


def normalize_deepseek_result(raw_result, fallback):
    if not isinstance(raw_result, dict):
        raise ValueError("DeepSeek returned a non-object JSON result")

    scenario = str(raw_result.get("scenario", "")).strip()
    print(f"Scenario: {scenario}")  # 调试信息
    if scenario not in SCENARIOS:
        raise ValueError("DeepSeek returned an unsupported scenario")

    normalized = {"scenario": scenario, "parse_method": "DeepSeek解析"}
    default_weights = RULE_PARSE_DEFAULTS[scenario]
    for field in WEIGHT_FIELDS:
        normalized[field] = coerce_weight(raw_result.get(field, default_weights[field]), default_weights[field])
    normalized = normalize_weights(normalized, default_weights)

    walking_time = raw_result.get("walking_time", fallback["walking_time"])
    try:
        normalized["walking_time"] = max(1, int(float(walking_time)))
    except (TypeError, ValueError):
        normalized["walking_time"] = fallback["walking_time"]
    explanation = raw_result.get("explanation") or fallback["explanation"]
    normalized["explanation"] = str(explanation).strip() or fallback["explanation"]
    return normalized


def call_deepseek_parser(user_text, selected_scene, fallback):
    api_key = get_config_value("DEEPSEEK_API_KEY")
    base_url = get_config_value("DEEPSEEK_BASE_URL", "https://api.sydney-ai.com/v1")
    model = get_config_value("DEEPSEEK_MODEL", "deepseek-chat")
    url = base_url.rstrip("/") + "/chat/completions"
    debug_info = {
        "是否读取到 DEEPSEEK_API_KEY": bool(api_key),
        "请求 URL": url,
        "模型名称": model,
        "HTTP 状态码": None,
        "DeepSeek 返回 content 前 300 字符": "",
        "JSON 解析失败原因": "",
    }
    st.session_state["deepseek_debug_info"] = debug_info

    if not api_key or not user_text.strip():
        if not api_key:
            debug_info["JSON 解析失败原因"] = "未调用：没有读取到 DEEPSEEK_API_KEY"
        else:
            debug_info["JSON 解析失败原因"] = "未调用：用户输入为空"
        return None

    prompt = f"""
请根据用户提供的社区规划需求生成一个标准 JSON 对象。
只返回 JSON，不要 Markdown，不要解释文字，不要代码块。

JSON 字段必须包含：
scenario, medical_weight, education_weight, commerce_weight, green_weight, traffic_weight, walking_time, explanation

要求：
1. scenario 只能是以下三个之一：候鸟老人友好模式、年轻家庭模式、游客短租模式。
2. 如果用户输入包含：老人、候鸟、养老、医院、买菜、慢行、休息座椅、公园、过马路少、康养，优先判断为：候鸟老人友好模式。
3. 如果用户输入包含：孩子、上学、学校、家庭、教育、运动、居住安静、亲子，优先判断为：年轻家庭模式。
4. 如果用户输入包含：游客、过来玩、旅游、大海、海边、短租、酒店、民宿、景点、免税店、公交接驳、夜间出行，优先判断为：游客短租模式。
5. 如果用户输入和当前页面选择场景冲突，以用户输入为准。当前页面选择场景只能作为默认参考，不要强制覆盖用户输入。
6. medical_weight、education_weight、commerce_weight、green_weight、traffic_weight 必须是 0 到 1 的小数。
7. walking_time 必须是整数分钟。
8. explanation 必须是中文一句话，简要说明解析依据。

当前页面选择场景：{selected_scene}
用户需求：{user_text}
"""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "你是社区规划需求解析器。只返回可被 json.loads 解析的 JSON 对象，不要 Markdown，不要解释文字。",
            },
            {"role": "user", "content": prompt},
        ],
        "stream": False,
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        debug_info["HTTP 状态码"] = response.status_code
        debug_info["DeepSeek 返回 content 前 300 字符"] = response.text[:300]
        response.raise_for_status()
        result = response.json()
        content = result["choices"][0]["message"]["content"].strip()
        debug_info["DeepSeek 返回 content 前 300 字符"] = content[:300]
        content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"\s*```$", "", content).strip()
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as e:
            debug_info["JSON 解析失败原因"] = str(e)
            print(f"Error parsing DeepSeek JSON: {str(e)}")
            return None
        return normalize_deepseek_result(parsed, fallback)
    except Exception as e:
        debug_info["JSON 解析失败原因"] = str(e)
        print(f"Error in DeepSeek API: {str(e)}")
        return None


def parse_user_need(user_text, selected_scene):
    fallback = parse_with_rules(user_text, selected_scene)
    try:
        deepseek_result = call_deepseek_parser(user_text, selected_scene, fallback)
        if deepseek_result:
            return deepseek_result
    except (ValueError, KeyError, TypeError, json.JSONDecodeError, error.HTTPError, error.URLError, TimeoutError, OSError):
        pass
    return fallback


st.set_page_config(
    page_title="琼岛智划",
    page_icon="🌴",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown(
    """
<style>
:root {
    --island-blue: #1b8fb8;
    --island-teal: #18a99d;
    --palm-green: #2f9b6d;
    --coral: #ff8f7a;
    --sand: #fff4df;
    --cream: #fffdf8;
    --sky: #eefaff;
    --ink: #17324d;
    --muted: #5f7285;
    --line: rgba(27, 143, 184, 0.16);
    --shadow: 0 12px 34px rgba(23, 50, 77, 0.10);
}

.stApp {
    background: linear-gradient(180deg, #f3fbff 0%, #fffdf8 48%, #fff7e8 100%);
    color: var(--ink);
}

[data-testid="stAppViewContainer"] > .main {
    background:
        linear-gradient(135deg, rgba(27, 143, 184, 0.08), rgba(47, 155, 109, 0.04) 38%, rgba(255, 143, 122, 0.07) 100%);
}

[data-testid="stHeader"] {
    background: rgba(243, 251, 255, 0.72);
    backdrop-filter: blur(10px);
}

.block-container {
    max-width: 1220px;
    padding-top: 1.3rem;
    padding-bottom: 2.5rem;
}

section[data-testid="stSidebar"] > div {
    background: linear-gradient(180deg, #e9fbff 0%, #fff9ed 100%);
    border-right: 1px solid rgba(27, 143, 184, 0.14);
}

section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] label {
    color: var(--ink);
}

h1, h2, h3 {
    color: var(--ink);
    letter-spacing: 0;
}

h2 {
    margin-top: 1.2rem;
    padding: 0.25rem 0 0.45rem;
    border-bottom: 1px solid var(--line);
}

.island-hero {
    padding: 1.55rem 1.75rem;
    margin: 0.25rem 0 1.25rem;
    border: 1px solid rgba(27, 143, 184, 0.18);
    border-radius: 24px;
    background: linear-gradient(135deg, rgba(255, 253, 248, 0.96), rgba(232, 249, 255, 0.96) 54%, rgba(255, 244, 223, 0.95));
    box-shadow: var(--shadow);
}

.hero-kicker {
    color: var(--island-teal);
    font-size: 0.92rem;
    font-weight: 700;
    margin-bottom: 0.35rem;
}

.hero-title {
    color: var(--ink);
    font-size: 2.35rem;
    line-height: 1.12;
    font-weight: 800;
    margin: 0;
}

.hero-subtitle {
    color: var(--muted);
    font-size: 1.02rem;
    line-height: 1.75;
    margin: 0.8rem 0 0;
}

.hero-pills {
    display: flex;
    flex-wrap: wrap;
    gap: 0.5rem;
    margin-top: 1rem;
}

.hero-pill {
    color: #0f5d73;
    background: rgba(255, 255, 255, 0.78);
    border: 1px solid rgba(27, 143, 184, 0.16);
    border-radius: 999px;
    padding: 0.35rem 0.72rem;
    font-size: 0.86rem;
    font-weight: 650;
}

.island-card {
    height: 100%;
    min-height: 158px;
    padding: 1.05rem 1.05rem 1rem;
    border-radius: 18px;
    background: rgba(255, 255, 255, 0.82);
    border: 1px solid rgba(27, 143, 184, 0.15);
    box-shadow: 0 10px 26px rgba(23, 50, 77, 0.08);
}

.island-card h3 {
    margin: 0 0 0.55rem;
    font-size: 1.05rem;
}

.island-card p {
    color: var(--muted);
    font-size: 0.94rem;
    line-height: 1.7;
    margin: 0;
}

.policy-card-grid {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    gap: 18px;
    align-items: stretch;
    margin-top: 0.75rem;
}

.policy-card {
    height: 100%;
    min-height: 190px;
    display: flex;
    flex-direction: column;
    justify-content: flex-start;
    padding: 1rem 1.05rem;
    border-radius: 18px;
    background: #ffffff;
    border: 1px solid #d7eef2;
    box-shadow: 0 8px 20px rgba(31, 53, 82, 0.06);
    box-sizing: border-box;
    min-width: 0;
}

.policy-card h3 {
    color: #16324f;
    font-size: 1.02rem;
    margin: 0 0 0.55rem;
}

.policy-card p {
    color: #50677f;
    line-height: 1.72;
    margin: 0;
    white-space: pre-wrap;
    overflow-wrap: anywhere;
}

@media (max-width: 900px) {
    .policy-card-grid {
        grid-template-columns: 1fr;
    }
}

.policy-source-card {
    background: #e9f7ef;
    border: 1px solid #bfe3cc;
    border-radius: 18px;
    color: #2e6b4b;
    padding: 0.82rem 1rem;
    margin: 0.35rem 0 0.85rem;
    box-shadow: 0 7px 16px rgba(46, 107, 75, 0.07);
}

.policy-source-card strong {
    color: #2e6b4b;
}

.policy-source-card ul {
    margin: 0.45rem 0 0 1.15rem;
    padding: 0;
}

.policy-source-card li {
    color: #2e6b4b;
    margin: 0.12rem 0;
}

.policy-evidence-grid {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    gap: 16px;
    margin: 0.75rem 0 1rem;
}

.policy-evidence-card {
    background: #fffdf8;
    border: 1px solid #d6eef2;
    border-radius: 18px;
    padding: 1rem 1.05rem;
    box-shadow: 0 8px 20px rgba(31, 53, 82, 0.055);
}

.policy-evidence-source {
    color: #16324f;
    font-weight: 800;
    margin-bottom: 0.55rem;
}

.policy-evidence-meta {
    color: #5f7285;
    font-size: 0.88rem;
    line-height: 1.55;
    margin: 0.12rem 0;
}

.policy-evidence-label {
    color: #0f5d73;
    font-size: 0.82rem;
    font-weight: 750;
    margin: 0.65rem 0 0.22rem;
}

.policy-evidence-snippet {
    color: #38566f;
    background: #f5fbff;
    border: 1px solid rgba(214, 238, 242, 0.82);
    border-radius: 14px;
    padding: 0.72rem 0.8rem;
    line-height: 1.72;
    margin: 0;
}

.policy-keyword-badges {
    display: flex;
    flex-wrap: wrap;
    gap: 0.38rem;
    margin: 0.25rem 0 0.1rem;
}

.policy-keyword-badge {
    color: #1f6f4f;
    background: #e8f8ef;
    border: 1px solid #bfe7cf;
    border-radius: 999px;
    padding: 0.14rem 0.5rem;
    font-size: 0.8rem;
    font-weight: 700;
}

.policy-evidence-reason {
    color: #50677f;
    line-height: 1.68;
    margin: 0;
}

.policy-evidence-empty {
    color: #61758a;
    background: #fffdf8;
    border: 1px solid #d6eef2;
    border-radius: 16px;
    padding: 0.82rem 0.95rem;
    margin: 0.65rem 0 0.9rem;
}

@media (max-width: 900px) {
    .policy-evidence-grid {
        grid-template-columns: 1fr;
    }
}

.result-note {
    color: #61758a;
    background: #fffdf8;
    border: 1px solid #d7eef2;
    border-radius: 16px;
    padding: 0.72rem 0.9rem;
    margin: 0.45rem 0 0.75rem;
    line-height: 1.7;
}

.result-caption {
    color: #61758a;
    text-align: center;
    font-size: 0.92rem;
    line-height: 1.65;
    margin-top: 0.65rem;
}

.land-use-legend-note {
    color: #0f5d73;
    background: #e9fbff;
    border: 1px solid rgba(27, 143, 184, 0.18);
    border-radius: 16px;
    padding: 0.72rem 0.9rem;
    margin: 0.35rem 0 0.75rem;
    line-height: 1.7;
}

.land-use-table-wrap {
    overflow-x: auto;
    border: 1px solid #d7eef2;
    border-radius: 16px;
    background: linear-gradient(180deg, #fffdf8 0%, #f5fbff 100%);
    box-shadow: 0 8px 20px rgba(31, 53, 82, 0.045);
    margin: 0.55rem 0 0.5rem;
}

.land-use-table {
    width: 100%;
    border-collapse: collapse;
    min-width: 720px;
}

.land-use-table th {
    color: #16324f;
    background: #e9f7ff;
    text-align: left;
    padding: 0.68rem 0.8rem;
    font-size: 0.92rem;
    font-weight: 750;
    border-bottom: 1px solid #d7eef2;
}

.land-use-table td {
    color: #435f78;
    padding: 0.62rem 0.8rem;
    border-bottom: 1px solid rgba(215, 238, 242, 0.75);
    line-height: 1.55;
    vertical-align: middle;
}

.land-use-table tbody tr:nth-child(even) {
    background: rgba(233, 247, 255, 0.42);
}

.land-use-table tbody tr:hover {
    background: rgba(232, 249, 239, 0.55);
}

.land-use-table tr:last-child td {
    border-bottom: 0;
}

.land-use-type {
    color: #16324f;
    font-weight: 700;
}

.legend-swatch {
    display: inline-block;
    width: 1.05rem;
    height: 1.05rem;
    border-radius: 4px;
    border: 1px solid rgba(31, 53, 82, 0.18);
    margin-right: 0.42rem;
    vertical-align: -0.18rem;
}

.legend-dot {
    border-radius: 999px;
}

.legend-color-cell {
    white-space: nowrap;
}

.legend-focus-cell {
    min-width: 5rem;
}

.legend-focus-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    padding: 0.14rem 0.52rem;
    border-radius: 999px;
    color: #1f6f4f;
    background: #e8f8ef;
    border: 1px solid #bfe7cf;
    font-size: 0.82rem;
    font-weight: 700;
}

.legend-focus-muted {
    color: #9badbd;
}

.summary-table-wrap {
    overflow-x: auto;
    max-width: 920px;
    margin: 0.65rem auto 0.35rem;
    border: 1px solid #d6eef2;
    border-radius: 16px;
    background: #fffdf8;
    box-shadow: 0 8px 18px rgba(31, 53, 82, 0.04);
}

.summary-table {
    width: 100%;
    border-collapse: collapse;
    min-width: 620px;
}

.summary-table th {
    color: #16324f;
    background: #eaf8fb;
    text-align: left;
    padding: 0.72rem 0.85rem;
    border-bottom: 1px solid #d6eef2;
    font-size: 0.93rem;
    font-weight: 750;
}

.summary-table td {
    color: #38566f;
    background: #ffffff;
    padding: 0.68rem 0.85rem;
    border-bottom: 1px solid rgba(214, 238, 242, 0.82);
    line-height: 1.55;
}

.summary-table tbody tr:nth-child(even) td {
    background: #fffdf8;
}

.summary-table tbody tr:hover td {
    background: #f0faf8;
}

.summary-table tr:last-child td {
    border-bottom: 0;
}

.planning-image-scroll {
    max-width: 680px;
    height: 540px;
    overflow-y: auto;
    overflow-x: hidden;
    margin: 0 auto;
    padding: 16px;
    background: #fffdf8;
    border: 1px solid #d6eef2;
    border-radius: 22px;
    box-shadow: 0 14px 36px rgba(36, 91, 112, 0.10);
}

.planning-image-scroll img {
    width: 100%;
    height: auto;
    display: block;
    border-radius: 14px;
}

.instant-suggestion-grid {
    display: grid;
    grid-template-columns: repeat(3, minmax(0, 1fr));
    gap: 0.85rem;
    margin: 0.8rem 0 0.55rem;
}

.instant-suggestion-card {
    min-height: 104px;
    padding: 0.88rem 0.95rem;
    border-radius: 16px;
    background: #fffdf8;
    border: 1px solid #d6eef2;
    box-shadow: 0 8px 18px rgba(31, 53, 82, 0.05);
}

.instant-suggestion-index {
    display: inline-flex;
    width: 1.65rem;
    height: 1.65rem;
    align-items: center;
    justify-content: center;
    border-radius: 999px;
    color: #ffffff;
    background: linear-gradient(135deg, #1b8fb8, #18a99d);
    font-weight: 800;
    margin-bottom: 0.48rem;
}

.instant-suggestion-card h4 {
    color: #16324f;
    margin: 0;
    font-size: 0.98rem;
    line-height: 1.45;
}

.instant-suggestion-card p {
    color: #61758a;
    margin: 0.35rem 0 0;
    font-size: 0.88rem;
    line-height: 1.55;
}

@media (max-width: 900px) {
    .instant-suggestion-grid {
        grid-template-columns: 1fr;
    }
}

.tech-flow-grid {
    display: grid;
    grid-template-columns: repeat(5, minmax(0, 1fr));
    gap: 0.85rem;
    margin: 0.85rem 0 0.75rem;
}

.tech-flow-card {
    min-height: 230px;
    padding: 1rem 0.95rem;
    border-radius: 18px;
    background: rgba(255, 253, 248, 0.94);
    border: 1px solid rgba(27, 143, 184, 0.16);
    box-shadow: 0 8px 20px rgba(31, 53, 82, 0.06);
}

.tech-flow-step {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 1.85rem;
    height: 1.85rem;
    border-radius: 999px;
    color: #ffffff;
    background: linear-gradient(135deg, var(--island-blue), var(--island-teal));
    font-weight: 800;
    margin-bottom: 0.55rem;
}

.tech-flow-card h3 {
    color: #16324f;
    font-size: 1rem;
    margin: 0 0 0.5rem;
}

.tech-flow-card ul {
    color: #50677f;
    margin: 0;
    padding-left: 1.05rem;
    line-height: 1.65;
    font-size: 0.9rem;
}

.tech-flow-card li {
    margin: 0.15rem 0;
}

.tech-flow-note {
    color: #0f5d73;
    background: #e9fbff;
    border: 1px solid rgba(27, 143, 184, 0.18);
    border-radius: 16px;
    padding: 0.82rem 0.95rem;
    margin: 0.45rem 0 1rem;
    line-height: 1.7;
}

@media (max-width: 1100px) {
    .tech-flow-grid {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }
}

@media (max-width: 700px) {
    .tech-flow-grid {
        grid-template-columns: 1fr;
    }
}

[data-testid="stVerticalBlockBorderWrapper"] {
    background: #ffffff !important;
    border: 1px solid #d7eef2 !important;
    border-radius: 18px !important;
    box-shadow: 0 8px 20px rgba(31, 53, 82, 0.06) !important;
}

[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid rgba(191, 220, 234, 0.85);
    border-radius: 18px;
    padding: 0.8rem 0.85rem;
    box-shadow: 0 8px 18px rgba(31, 53, 82, 0.07);
}

[data-testid="stMetricLabel"] {
    color: #61758a !important;
}

[data-testid="stMetricValue"] {
    color: #2d8db3 !important;
}

.stTextArea textarea {
    border-radius: 18px;
    border: 1px solid #bfdcea !important;
    background: #fffdf8 !important;
    color: #1f3552 !important;
    caret-color: var(--island-blue);
    box-shadow: 0 5px 14px rgba(31, 53, 82, 0.045);
}

.stTextArea textarea::placeholder {
    color: #7a8ca5 !important;
    opacity: 1 !important;
}

.stTextArea textarea:focus {
    border-color: #5ab7d6 !important;
    box-shadow: 0 0 0 3px rgba(90, 183, 214, 0.16), 0 5px 14px rgba(31, 53, 82, 0.045) !important;
}

div[data-baseweb="select"] > div,
[role="radiogroup"] {
    border-radius: 16px;
    border-color: rgba(27, 143, 184, 0.18);
}

div[data-baseweb="select"] > div {
    background: #fffdf8 !important;
    border: 1px solid #bfdcea !important;
    color: #1f3552 !important;
    box-shadow: 0 5px 14px rgba(31, 53, 82, 0.045);
}

div[data-baseweb="select"] span,
div[data-baseweb="select"] input,
div[data-baseweb="select"] svg {
    color: #1f3552 !important;
    fill: #2d8db3 !important;
}

div[data-baseweb="popover"],
div[data-baseweb="menu"],
ul[role="listbox"] {
    background: #fffdf8 !important;
    color: #1f3552 !important;
    border: 1px solid #bfdcea !important;
    border-radius: 16px !important;
    box-shadow: 0 14px 28px rgba(31, 53, 82, 0.12) !important;
}

div[data-baseweb="popover"] div,
div[data-baseweb="menu"] div,
ul[role="listbox"] div {
    background-color: #fffdf8 !important;
    color: #1f3552 !important;
}

li[role="option"],
div[role="option"] {
    background: #fffdf8 !important;
    color: #1f3552 !important;
}

li[role="option"]:hover,
div[role="option"]:hover,
li[aria-selected="true"],
div[aria-selected="true"] {
    background: #e9fbff !important;
    color: #0f5d73 !important;
}

li[role="option"]:hover *,
div[role="option"]:hover *,
li[aria-selected="true"] *,
div[aria-selected="true"] * {
    background: #e9fbff !important;
    color: #0f5d73 !important;
}

[data-testid="stRadio"] label,
[data-testid="stCheckbox"] label,
[data-testid="stRadio"] p,
[data-testid="stCheckbox"] p {
    color: #1f3552 !important;
}

[data-baseweb="radio"],
[data-baseweb="checkbox"] {
    color: #1f3552 !important;
}

[data-baseweb="radio"] > div:first-child,
[data-baseweb="checkbox"] > div:first-child {
    border-color: #5ab7d6 !important;
    background-color: #fffdf8 !important;
}

[data-baseweb="radio"][aria-checked="true"] > div:first-child,
[data-baseweb="checkbox"][aria-checked="true"] > div:first-child {
    border-color: #18a99d !important;
    background-color: #18a99d !important;
}

.status-card {
    border-radius: 18px;
    padding: 0.78rem 0.92rem;
    margin: 0 0 0.65rem;
    border: 1px solid rgba(31, 53, 82, 0.08);
    box-shadow: 0 7px 16px rgba(31, 53, 82, 0.055);
}

.status-label {
    font-size: 0.78rem;
    font-weight: 750;
    margin-bottom: 0.18rem;
}

.status-value {
    font-size: 0.98rem;
    font-weight: 750;
    line-height: 1.45;
}

.status-scene {
    background: #ddf4e4;
    color: #246b45;
}

.status-parse {
    background: #ddeeff;
    color: #23527c;
}

.status-auto {
    background: #fff4cc;
    color: #7a5600;
}

.status-scene .status-label,
.status-scene .status-value {
    color: #246b45;
}

.status-parse .status-label,
.status-parse .status-value {
    color: #23527c;
}

.status-auto .status-label,
.status-auto .status-value {
    color: #7a5600;
}

[data-testid="stVegaLiteChart"],
[data-testid="stArrowVegaLiteChart"] {
    background: #ffffff !important;
    border: 1px solid #d7eef2;
    border-radius: 18px;
    padding: 0.8rem;
    box-shadow: 0 8px 18px rgba(31, 53, 82, 0.06);
}

[data-testid="stVegaLiteChart"] canvas,
[data-testid="stArrowVegaLiteChart"] canvas,
[data-testid="stVegaLiteChart"] svg,
[data-testid="stArrowVegaLiteChart"] svg {
    background: #ffffff !important;
}

.stAlert {
    border-radius: 18px;
    border: 1px solid rgba(27, 143, 184, 0.14);
    box-shadow: 0 8px 22px rgba(23, 50, 77, 0.06);
}

[data-testid="stExpander"],
[data-testid="stDataFrame"],
[data-testid="stImage"] img {
    border-radius: 18px;
}

[data-testid="stDataFrame"] {
    background: #ffffff !important;
    border: 1px solid #d7eef2;
    box-shadow: 0 8px 18px rgba(31, 53, 82, 0.05);
}

[data-testid="stImage"] img {
    box-shadow: var(--shadow);
}

.stTabs [data-baseweb="tab-list"] {
    gap: 0.45rem;
    padding: 0.35rem;
    border-radius: 18px;
    background: rgba(233, 251, 255, 0.75);
    border: 1px solid rgba(27, 143, 184, 0.12);
}

.stTabs [data-baseweb="tab"] {
    border-radius: 14px;
    color: var(--muted);
    padding: 0.55rem 0.9rem;
}

.stTabs [aria-selected="true"] {
    color: #0f5d73;
    background: rgba(255, 255, 255, 0.9);
    box-shadow: 0 6px 16px rgba(23, 50, 77, 0.08);
}

.stDownloadButton button {
    border: 0;
    border-radius: 999px;
    color: white;
    background: linear-gradient(135deg, var(--island-blue), var(--island-teal));
    box-shadow: 0 10px 24px rgba(24, 169, 157, 0.26);
    padding: 0.7rem 1.15rem;
    font-weight: 700;
}

.stDownloadButton button:hover {
    border: 0;
    color: white;
    filter: brightness(1.03);
}

hr {
    border-color: rgba(27, 143, 184, 0.12);
    margin: 1.25rem 0;
}
</style>
""",
    unsafe_allow_html=True,
)

st.sidebar.title("🌴 琼岛智划")
st.sidebar.caption("面向海南自贸港的 AI 多智能体社区规划系统")

scene = st.sidebar.radio(
    "选择规划场景",
    SCENARIOS,
    index=0
)

st.sidebar.divider()

show_debug_info = st.sidebar.checkbox("显示开发调试信息", value=False)

st.sidebar.divider()

st.sidebar.subheader("系统模块")
st.sidebar.write("✅ 离线规划引擎 + 在线智能解释")
st.sidebar.write("✅ 上游 PPO-GNN/SGNN 底层依据")
st.sidebar.write("✅ 自然语言需求解析")
st.sidebar.write("✅ 多Agent解释辅助")
st.sidebar.write("✅ RAG合规辅助")
st.sidebar.write("✅ 规划图与指标展示")

st.sidebar.divider()

st.sidebar.info(
    "当前演示端读取上游 PPO/SGNN 项目离线生成的规划结果图和统计表，不进行现场训练。"
)

st.markdown(
    """
<div class="island-hero">
  <div class="hero-kicker">🌊 Hainan Free Trade Port · Sanya Island Planning</div>
  <h1 class="hero-title">🌴 琼岛智划</h1>
  <p class="hero-subtitle">面向海南自贸港的多智能体协同社区规划系统，把自然语言需求转化为清晰、可解释、可展示的 15 分钟生活圈优化方案。</p>
  <div class="hero-pills">
    <span class="hero-pill">海风感交互</span>
    <span class="hero-pill">三亚场景适配</span>
    <span class="hero-pill">离线规划引擎 + 在线智能解释</span>
    <span class="hero-pill">政策可解释</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
本系统面向海南自贸港建设背景下的社区规划需求，采用 **离线规划引擎 + 在线智能解释** 架构：
展示端读取上游 **PPO/SGNN** 项目离线生成的规划结果图和统计表，不进行现场训练；
DeepSeek、多Agent、RAG 模块用于需求解析、解释生成和合规辅助，不直接改变底层空间规划结果。
"""
)

st.markdown("## 海南特色场景适配")
st.markdown("针对海南自贸港建设与三亚等热带滨海城市特点，系统提供三大特色场景适配方案：")

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown(
        """
<div class="island-card">
  <h3>🏠 候鸟老人旅居养老</h3>
  <p>冬季人口潮汐流动明显，系统重点关注基层医疗、菜市场、公园绿地和慢行交通，提高旅居老人日常生活便利性。</p>
</div>
""",
        unsafe_allow_html=True,
    )

with col2:
    st.markdown(
        """
<div class="island-card">
  <h3>🛍️ 自贸港旅游消费</h3>
  <p>结合海南自贸港和国际旅游消费中心建设需求，系统关注酒店民宿、免税商业、景点连接和交通接驳。</p>
</div>
""",
        unsafe_allow_html=True,
    )

with col3:
    st.markdown(
        """
<div class="island-card">
  <h3>🌴 热带滨海社区治理</h3>
  <p>面向三亚等热带滨海城市，系统兼顾生态绿地、公共服务均衡、道路连通和社区生活秩序。</p>
</div>
""",
        unsafe_allow_html=True,
    )

st.divider()

st.markdown("## 系统工作流")

flow_steps = [
    {
        "icon": "📝",
        "title": "自然语言需求输入",
        "desc": "用户以自然语言描述社区规划需求，支持手动输入或选择演示样例。"
    },
    {
        "icon": "🔍",
        "title": "场景识别与指标权重",
        "desc": "系统根据关键词自动识别场景（候鸟老人/年轻家庭/游客短租），并生成对应指标权重配置。"
    },
    {
        "icon": "🤝",
        "title": "多Agent解释辅助",
        "desc": "居民Agent、政府Agent、产业运营Agent模拟多方诉求，协调器Agent输出解释性折中建议。"
    },
    {
        "icon": "🗺️",
        "title": "上游PPO-GNN/SGNN离线结果",
        "desc": "读取上游 PPO-GNN/SGNN 项目离线生成的规划结果图与空间指标统计表。"
    },
    {
        "icon": "📚",
        "title": "RAG合规辅助",
        "desc": "基于本地政策文本检索增强，生成场景特定的合规性解释文本。"
    },
    {
        "icon": "📋",
        "title": "一键生成报告",
        "desc": "自动整合所有内容，生成可下载的Markdown格式规划分析报告。"
    }
]

cols = st.columns(len(flow_steps))
for i, step in enumerate(flow_steps):
    with cols[i]:
        st.markdown(f"### {step['icon']} {step['title']}")
        st.write(step['desc'])

st.divider()

demo_samples = {
    "候鸟老人社区": "这个社区冬季候鸟老人较多，希望步行10分钟内可以到社区医院、菜市场和公园，路上需要休息空间，过马路尽量少。",
    "年轻家庭社区": "这个片区年轻家庭较多，希望孩子上学方便，周边有商业配套和运动休闲空间，同时保证居住环境安静。",
    "游客短租社区": "这个区域民宿和酒店较多，希望游客从公交站、景点、商业街和免税消费点之间换乘方便，夜间出行也要安全。"
}

left, right = st.columns([2, 1])

with left:
    selected_sample = st.selectbox(
        "选择演示样例",
        ["不使用样例", "候鸟老人社区", "年轻家庭社区", "游客短租社区"],
        index=0,
        help="选择一个演示样例快速填充输入框"
    )
    
    default_input = demo_samples.get(selected_sample, "")
    
    user_input = st.text_area(
        "请输入你的规划需求",
        value=default_input,
        placeholder="例如：这个社区候鸟老人比较多，希望步行10分钟内能到医院、菜市场和公园，绿地多一点，过马路少一点。",
        height=120
    )

parse_result = parse_user_need(user_input, scene)
auto_scene = parse_result["scenario"]

with right:
    st.markdown("### 当前选择场景")
    st.markdown(
        f"""
<div class="status-card status-scene">
  <div class="status-label">当前选择场景</div>
  <div class="status-value">{html.escape(scene)}</div>
</div>
""",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"""
<div class="status-card status-parse">
  <div class="status-label">解析方式</div>
  <div class="status-value">{html.escape(parse_result['parse_method'])}</div>
</div>
""",
        unsafe_allow_html=True,
    )
    
    if auto_scene != scene:
        scene_note = f"系统自动识别为：{auto_scene}"
    else:
        scene_note = f"系统采用：{scene}"

    st.markdown(
        f"""
<div class="status-card status-auto">
  <div class="status-label">系统识别说明</div>
  <div class="status-value">{html.escape(scene_note)}</div>
</div>
""",
        unsafe_allow_html=True,
    )

    st.caption(parse_result["explanation"])

    if show_debug_info:
        with st.expander("DeepSeek 调试信息"):
            st.json(st.session_state.get("deepseek_debug_info", {}))

final_scene = auto_scene

hlg_image = HLG_IMG if HLG_IMG.exists() else DHM_IMG
hlg_summary = HLG_SUMMARY if HLG_SUMMARY.exists() else DHM_SUMMARY

scene_configs = {
    "候鸟老人友好模式": {
        "image": DHM_IMG,
        "summary": DHM_SUMMARY,
        "geojson": DHM_GEOJSON,
        "main_goal": "提升医疗、菜市场、公园绿地和慢行交通的可达性",
        "weights": {"医疗可达性": 0.35, "菜市场/商业便利": 0.25, "公园绿地": 0.25, "道路连通": 0.15},
        "recommendation": "优先采用 DHM 社区优化结果，用于展示候鸟老人生活圈友好布局。"
    },
    "年轻家庭模式": {
        "image": hlg_image,
        "summary": hlg_summary,
        "geojson": HLG_GEOJSON,
        "main_goal": "提升教育、商业、就业和社区活力设施的可达性",
        "weights": {"教育可达性": 0.35, "商业服务": 0.25, "就业/办公": 0.20, "绿地休闲": 0.20},
        "recommendation": "优先展示 HLG 或 DHM 规划结果，并从教育与家庭生活便利角度解释。"
    },
    "游客短租模式": {
        "image": hlg_image,
        "summary": hlg_summary,
        "geojson": HLG_GEOJSON,
        "main_goal": "提升交通接驳、商业消费、景点连接和短住服务便利性",
        "weights": {"交通接驳": 0.35, "商业/免税消费": 0.30, "景点可达": 0.20, "公共空间": 0.15},
        "recommendation": "优先从交通和商业服务角度展示规划结果，适合旅游短租社区解释。"
    }
}

config = scene_configs.get(final_scene, scene_configs["候鸟老人友好模式"])
geojson_path = config.get("geojson", DHM_GEOJSON)
focus_types = SCENE_FOCUS_TYPES.get(final_scene, set())

st.markdown("## 1. 场景解析与指标权重")
st.caption(f"解析方式：{parse_result['parse_method']}｜建议步行时间：{parse_result['walking_time']} 分钟")

parsed_weight_items = [
    ("医疗权重", parse_result["medical_weight"]),
    ("教育权重", parse_result["education_weight"]),
    ("商业权重", parse_result["commerce_weight"]),
    ("绿地权重", parse_result["green_weight"]),
    ("交通权重", parse_result["traffic_weight"]),
]

for col, (label, value) in zip(st.columns(len(parsed_weight_items)), parsed_weight_items):
    with col:
        st.metric(label, f"{value:.0%}")

with st.expander("查看需求解析 JSON"):
    st.json(
        {
            "scenario": parse_result["scenario"],
            "medical_weight": parse_result["medical_weight"],
            "education_weight": parse_result["education_weight"],
            "commerce_weight": parse_result["commerce_weight"],
            "green_weight": parse_result["green_weight"],
            "traffic_weight": parse_result["traffic_weight"],
            "walking_time": parse_result["walking_time"],
            "explanation": parse_result["explanation"],
        }
    )

col_left, col_right = st.columns([1, 1])
with col_left:
    st.markdown("### 核心规划目标")
    st.write(config["main_goal"])
with col_right:
    st.markdown("### 推荐方案")
    st.write(config["recommendation"])

st.markdown("## 2. 规划结果可视化")
st.markdown(
    """
<div class="result-note">
系统读取已生成的空间规划结果和统计指标，用于展示当前场景下的规划空间结构与服务设施分布。
</div>
""",
    unsafe_allow_html=True,
)

img_path = config["image"]
summary_path = config["summary"]
summary_df = None
numeric_cols = []
summary_error = None

if summary_path.exists():
    try:
        summary_df = pd.read_csv(summary_path)
        numeric_cols = summary_df.select_dtypes(include=["number"]).columns.tolist()
    except Exception as e:
        summary_error = str(e)

result_tab_image, result_tab_geojson, result_tab_instant, result_tab_table, result_tab_chart = st.tabs(["规划结果图", "动态GeoJSON图", "即时建议图", "空间统计表", "数值指标图"])

with result_tab_image:
    with st.container(border=True):
        if img_path.exists():
            if img_path.stat().st_size > 0:
                try:
                    if PIL_AVAILABLE:
                        with Image.open(img_path) as image:
                            image.verify()
                    render_planning_image_scroll(img_path)
                    st.markdown(
                        """
<div class="result-caption">
该图为上游 PPO/SGNN 项目离线生成结果的展示图，当前演示端不进行现场训练或实时生成新规划图。
</div>
""",
                        unsafe_allow_html=True,
                    )
                except Exception as e:
                    st.warning("图片文件存在，但暂时无法识别为有效图片。")
                    st.caption(f"图片路径：{img_path}｜错误原因：{e}")
            else:
                st.warning("图片文件为空，请重新生成规划图。")
                st.caption(f"当前图片路径：{img_path}")
        else:
            st.warning(f"没有找到规划图：{img_path}")
            st.caption("请确认 outputs 文件夹中存在 dhm_result_clean.png 和 hlg_result_clean.png。")

    with st.container(border=True):
        st.markdown("### 用地类型图例说明")
        st.caption("type 含义来自上游 DRL urban planning 项目的 `city_config.py`。当前页面图片为展示端基于 GeoJSON `type` 字段生成的可视化结果，颜色采用展示端渲染配色。")
        st.markdown(render_land_use_legend_html(), unsafe_allow_html=True)
        st.caption("注：颜色为展示端渲染配色，土地类型含义以上游 `city_config.py` 的 `type` 定义为准。其中 `type=4` 表示**居住用地**，`type=7`、`type=8` 表示**绿地类用地**。")

with result_tab_geojson:
    with st.container(border=True):
        st.markdown("### 动态GeoJSON图")
        st.caption("该图为展示端基于离线 GeoJSON 结果的实时渲染，不代表现场重新训练 PPO/SGNN。")
        st.caption("当前场景重点类型通过边框、线宽和图例标注进行提示，不改变底层空间规划结果。")

        try:
            fig, feature_count = render_geojson_plan_figure(geojson_path, focus_types)
            geojson_image = render_matplotlib_figure_png(fig)
            render_centered_image(geojson_image, caption=f"{final_scene} - 动态GeoJSON图", width=520)
            if MATPLOTLIB_AVAILABLE:
                plt.close(fig)
            st.caption(f"已读取：{geojson_path.name}｜GeoJSON 要素数：{feature_count}")
        except Exception as e:
            st.warning("未找到有效 GeoJSON，当前显示静态规划结果图。")
            st.caption(f"GeoJSON 路径：{geojson_path}｜错误原因：{e}")

    with st.container(border=True):
        st.markdown("### 动态图图例")
        st.caption("type 含义来自上游 DRL urban planning 项目的 `city_config.py`。当前页面图片为展示端基于 GeoJSON `type` 字段生成的可视化结果，颜色采用展示端渲染配色。")
        st.markdown(render_land_use_legend_html(focus_types, include_focus=True), unsafe_allow_html=True)
        st.caption("注：颜色为展示端渲染配色，土地类型含义以上游 `city_config.py` 的 `type` 定义为准。其中 `type=4` 表示**居住用地**，`type=7`、`type=8` 表示**绿地类用地**。")

with result_tab_instant:
    with st.container(border=True):
        st.markdown("### 即时规划建议图")
        st.caption("该图为展示端基于离线 GeoJSON 和当前场景生成的即时标注图，图中编号对应下方三条建议。")

        try:
            instant_fig, instant_config = render_instant_plan_figure(geojson_path, final_scene, parse_result)
            instant_image = render_matplotlib_figure_png(instant_fig)
            render_centered_image(instant_image, caption=f"{final_scene} - 即时规划建议图", width=560)
            if MATPLOTLIB_AVAILABLE:
                plt.close(instant_fig)

            st.download_button(
                "下载即时建议图 PNG",
                data=instant_image,
                file_name=f"琼岛智划_{final_scene}_即时建议图.png",
                mime="image/png",
            )

            focus_type_names = "、".join(get_focus_type_names(instant_config["focus_types"]))
            st.caption(f"重点关注类型：{focus_type_names}")
            render_instant_suggestion_cards(instant_config["suggestions"])
            st.caption("该图为基于离线 GeoJSON 的即时场景化标注图，用于解释用户需求下的优化方向；当前版本不进行现场 PPO/SGNN 训练，也不实时生成新的底层规划结果。")
        except Exception as e:
            st.warning("暂时无法生成即时建议图，当前仍可查看静态规划结果图和动态 GeoJSON 图。")
            st.caption(f"GeoJSON 路径：{geojson_path}｜错误原因：{e}")

with result_tab_table:
    with st.container(border=True):
        st.markdown(
            """
<div class="result-note">
以下统计表展示规划地块面积、道路长度、设施类型等空间指标。
</div>
""",
            unsafe_allow_html=True,
        )
        if summary_df is not None:
            st.markdown(render_summary_table_html(summary_df), unsafe_allow_html=True)
        elif summary_error:
            st.warning(f"统计表读取失败：{summary_error}")
        else:
            st.warning(f"没有找到统计表：{summary_path}")
            st.caption("请确认 data 文件夹中存在对应的 CSV 文件。")

with result_tab_chart:
    with st.container(border=True):
        if summary_df is not None and numeric_cols:
            st.bar_chart(summary_df[numeric_cols], height=320)
            st.markdown(
                """
<div class="result-caption">
数值图用于辅助观察不同空间指标的相对变化。
</div>
""",
                unsafe_allow_html=True,
            )
        elif summary_df is not None:
            st.info("当前统计表中没有可用于绘图的数值字段。")
        elif summary_error:
            st.warning(f"数值图生成失败：{summary_error}")
        else:
            st.warning("暂时无法生成数值指标图，因为未找到对应统计表。")

st.markdown("## 3. 离线规划引擎 + 在线智能解释")
st.markdown(
    """
<div class="tech-flow-note">
当前演示端读取上游 PPO/SGNN 项目离线生成的规划结果图和统计表，不进行现场训练；DeepSeek、多Agent、RAG 模块用于需求解析、解释生成和合规辅助，不直接改变底层空间规划结果。
</div>
<div class="tech-flow-grid">
  <div class="tech-flow-card">
    <div class="tech-flow-step">1</div>
    <h3>上游输入数据</h3>
    <ul>
      <li>社区空间地块数据</li>
      <li>道路、设施、用地类型等空间属性</li>
    </ul>
  </div>
  <div class="tech-flow-card">
    <div class="tech-flow-step">2</div>
    <h3>图结构建模</h3>
    <ul>
      <li>将地块或空间单元抽象为节点</li>
      <li>邻接关系、道路连接或空间相邻关系作为边</li>
    </ul>
  </div>
  <div class="tech-flow-card">
    <div class="tech-flow-step">3</div>
    <h3>多目标奖励函数</h3>
    <ul>
      <li>医疗可达</li>
      <li>教育可达</li>
      <li>商业便利</li>
      <li>绿地覆盖</li>
      <li>交通连通</li>
    </ul>
  </div>
  <div class="tech-flow-card">
    <div class="tech-flow-step">4</div>
    <h3>离线结果读取</h3>
    <ul>
      <li>当前演示端不现场训练</li>
      <li>不实时推理生成新规划图</li>
      <li>读取上游离线规划结果图和统计表</li>
      <li>保证比赛现场演示稳定</li>
    </ul>
  </div>
  <div class="tech-flow-card">
    <div class="tech-flow-step">5</div>
    <h3>在线智能解释</h3>
    <ul>
      <li>规划结果图</li>
      <li>空间统计表</li>
      <li>需求解析、合规解释与报告</li>
    </ul>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown("## 4. 多Agent协同分析")

AGENT_ORDER = [
    ("resident_agent", "居民Agent"),
    ("government_agent", "政府Agent"),
    ("business_agent", "商业Agent"),
    ("coordinator_agent", "协调器Agent"),
]

REPORT_AGENT_NAMES = {
    "resident_agent": "居民代表",
    "government_agent": "政府代理人",
    "business_agent": "商业代理",
    "coordinator_agent": "协调器Agent",
}

AGENT_FIELD_ALIASES = {
    "core_need": ["core_need", "核心诉求", "core_demand"],
    "reason": ["reason", "支持理由", "support_reason"],
    "conflict": ["conflict", "潜在冲突", "potential_conflict"],
    "compromise": ["compromise", "妥协建议", "compromise_suggestion", "suggestion"],
}


def build_rule_agent_logs(user_text, parse_result):
    scene_name = parse_result["scenario"]
    walking_time = parse_result["walking_time"]
    need_text = user_text.strip() or "用户暂未输入具体需求，系统根据默认场景进行规划解释。"

    if scene_name == "候鸟老人友好模式":
        return {
            "resident_agent": {
                "core_need": f"结合用户需求「{need_text}」，优先保障医院、买菜、公园、慢行、休息座椅和过马路少。",
                "reason": f"候鸟老人日常活动半径更依赖步行，医疗、菜场和公园应尽量控制在 {walking_time} 分钟可达范围内。",
                "conflict": "老人安静、安全和低干扰的生活需求，可能与商业客流、游客动线和道路改造强度产生冲突。",
                "compromise": "把医疗、菜场、公园和休息座椅布置在居住区近端，商业服务适度靠近但避免穿越老人主要慢行路径。",
            },
            "government_agent": {
                "core_need": "围绕养老服务、15分钟生活圈、公共服务均衡和慢行安全组织社区更新。",
                "reason": "该场景需要兼顾适老化、公共服务公平、交通组织和政策合规，避免只优化单一设施点。",
                "conflict": "公共资源均衡配置可能限制局部高强度商业开发，也可能与居民对近距离设施的强需求产生张力。",
                "compromise": "以15分钟生活圈为底线，优先补齐基层医疗、慢行过街安全和公共绿地，再安排便民商业。",
            },
            "business_agent": {
                "core_need": "发展康养消费、社区商业、药店、菜场和便民服务。",
                "reason": "适度商业可以支撑老人日常消费和康养服务，也能提升社区运营活力。",
                "conflict": "商业噪声、装卸和客流可能影响老人休息与慢行安全。",
                "compromise": "将药店、菜场和便民服务控制在小尺度、低噪声形态，并与主要休息空间保持缓冲。",
            },
            "coordinator_agent": {
                "core_need": "综合居民、政府、商业三方目标，以老人便利和慢行安全为优先。",
                "reason": "候鸟老人友好模式必须先满足公共服务、日常生活便利和安全步行，再用商业活力补足服务供给。",
                "conflict": "主要冲突在老人安静生活、商业服务便利和政府公共资源均衡之间。",
                "compromise": "采用近端公共服务、边缘适度商业、慢行优先交通组织和共享绿地空间的折中方案，控制商业干扰。",
            },
        }
    elif scene_name == "年轻家庭模式":
        return {
            "resident_agent": {
                "core_need": f"结合用户需求「{need_text}」，优先保障学校、亲子空间、运动场和居住安静。",
                "reason": f"年轻家庭对上学便利、亲子活动和日常生活配套敏感，核心设施应尽量在 {walking_time} 分钟范围内。",
                "conflict": "亲子活动和社区商业会带来人流，可能影响居住安静和交通安全。",
                "compromise": "把学校、运动场和亲子空间放在安全慢行路径上，商业布置在外围节点，减少对住宅的噪声影响。",
            },
            "government_agent": {
                "core_need": "优化教育资源、公共空间、交通安全和社区治理。",
                "reason": "年轻家庭模式需要保障教育公平、儿童友好空间和安全通学，同时维持社区运行秩序。",
                "conflict": "教育设施和公共空间集中可能带来上下学拥堵，也可能挤压其他公共服务用地。",
                "compromise": "通过分散公共空间、优化接送交通组织和设置安全步行线路，平衡教育可达与社区治理。",
            },
            "business_agent": {
                "core_need": "发展亲子消费、社区商业、培训/文体服务和便利生活配套。",
                "reason": "家庭型消费可以提升社区活力，并为亲子、运动和日常采购提供补充服务。",
                "conflict": "培训和商业业态可能加剧交通、噪声和停车压力。",
                "compromise": "将亲子商业和文体服务集中在社区边缘或公共交通节点，控制营业时段和噪声外溢。",
            },
            "coordinator_agent": {
                "core_need": "综合居民、政府、商业三方目标，优先保障教育与生活配套，同时保持居住安静。",
                "reason": "年轻家庭模式需要同时平衡公共服务、家庭生活便利和适度商业活力。",
                "conflict": "主要冲突在儿童活动强度、商业人流和住宅安静环境之间。",
                "compromise": "采用学校与亲子空间优先、商业外围布局、分时活动管理和安全交通组织的折中方案。",
            },
        }
    else:
        return {
            "resident_agent": {
                "core_need": f"结合用户需求「{need_text}」，重点控制噪声、维护生活秩序，避免公共资源被游客过度挤占。",
                "reason": "游客短租会带来流动人口和夜间活动，原住居民更关注安静、安全和日常生活不被干扰。",
                "conflict": "游客消费、夜间经济和景点接驳可能与居民休息、停车和公共空间使用发生冲突。",
                "compromise": "将游客动线与居民日常动线适度分离，对夜间活动和高峰客流进行时段管理。",
            },
            "government_agent": {
                "core_need": "强化旅游治理、交通接驳、公共安全和秩序管理。",
                "reason": "游客短租模式需要保障旅游便利，也要避免交通拥堵、公共空间超载和社区治理失序。",
                "conflict": "提升旅游接待能力可能增加社区管理成本，并对居民生活秩序造成压力。",
                "compromise": "通过公交接驳、步行导流、夜间安全管理和公共空间承载控制平衡游客与居民需求。",
            },
            "business_agent": {
                "core_need": "发展民宿酒店、免税商业、景点接驳和夜间经济。",
                "reason": "游客消费和交通接驳是该场景的核心活力来源，可提升片区旅游服务品质。",
                "conflict": "高强度商业和夜间经济可能带来噪声、拥堵和治理压力。",
                "compromise": "把民宿酒店、免税商业和夜间消费集中到适合承载的界面，避开居民安静生活空间。",
            },
            "coordinator_agent": {
                "core_need": "综合居民、政府、商业三方目标，在游客便利和居民生活之间做分区与时段协调。",
                "reason": "游客短租模式必须平衡公共服务、生活便利和商业活力，否则容易造成游客体验和居民秩序双重受损。",
                "conflict": "主要冲突在居民安静生活、游客消费需求和政府治理承载之间。",
                "compromise": "采用游客服务分区、夜间时段管理、交通接驳优化和公共空间共享规则，既方便游客也保护居民生活。",
            },
        }


def pick_agent_field(agent_value, field_name, fallback_value):
    if isinstance(agent_value, dict):
        for alias in AGENT_FIELD_ALIASES[field_name]:
            value = agent_value.get(alias)
            if value is not None and str(value).strip():
                return str(value).strip()
    elif field_name == "core_need" and agent_value is not None and str(agent_value).strip():
        return str(agent_value).strip()
    return fallback_value


def normalize_agent_logs(raw_result, fallback_logs):
    if not isinstance(raw_result, dict):
        return fallback_logs

    normalized = {}
    for agent_key, _ in AGENT_ORDER:
        agent_value = raw_result.get(agent_key, {})
        fallback_agent = fallback_logs[agent_key]
        normalized[agent_key] = {
            "core_need": pick_agent_field(agent_value, "core_need", fallback_agent["core_need"]),
            "reason": pick_agent_field(agent_value, "reason", fallback_agent["reason"]),
            "conflict": pick_agent_field(agent_value, "conflict", fallback_agent["conflict"]),
            "compromise": pick_agent_field(agent_value, "compromise", fallback_agent["compromise"]),
        }
    return normalized


def render_agent_content(agent_data):
    return (
        f"**核心诉求：** {agent_data['core_need']}\n\n"
        f"**支持理由：** {agent_data['reason']}\n\n"
        f"**潜在冲突：** {agent_data['conflict']}\n\n"
        f"**妥协建议：** {agent_data['compromise']}"
    )


def generate_agent_logs(user_text, parse_result):
    fallback_logs = build_rule_agent_logs(user_text, parse_result)
    api_key = get_config_value("DEEPSEEK_API_KEY")
    base_url = get_config_value("DEEPSEEK_BASE_URL", "https://api.sydney-ai.com/v1")
    model = get_config_value("DEEPSEEK_MODEL", "deepseek-chat")
    if not api_key:
        return fallback_logs

    need_text = user_text.strip() or "用户暂未输入具体需求，系统根据默认场景进行规划解释。"
    url = base_url.rstrip("/") + "/chat/completions"
    prompt = f"""
请基于用户需求和解析结果，生成三 Agent 协同过程 JSON。
只返回 JSON，不要 Markdown，不要解释文字，不要代码块。

JSON 必须严格使用以下结构，每个字段都不能为空：
{{
  "resident_agent": {{"core_need": "...", "reason": "...", "conflict": "...", "compromise": "..."}},
  "government_agent": {{"core_need": "...", "reason": "...", "conflict": "...", "compromise": "..."}},
  "business_agent": {{"core_need": "...", "reason": "...", "conflict": "...", "compromise": "..."}},
  "coordinator_agent": {{"core_need": "...", "reason": "...", "conflict": "...", "compromise": "..."}}
}}

请紧扣用户输入和 scenario：
- 候鸟老人友好模式：居民关注医院、买菜、公园、慢行、休息座椅、过马路少；政府关注养老服务、15分钟生活圈、公共服务均衡、慢行安全；商业关注康养消费、社区商业、药店、菜场、便民服务；协调器强调老人便利优先，同时控制商业干扰。
- 年轻家庭模式：居民关注学校、亲子空间、运动场、居住安静；政府关注教育资源、公共空间、交通安全、社区治理；商业关注亲子消费、社区商业、培训/文体服务；协调器强调教育与生活配套优先，同时保持居住安静。
- 游客短租模式：居民关注噪声控制、生活秩序、公共资源不被挤占；政府关注旅游治理、交通接驳、公共安全、秩序管理；商业关注民宿酒店、免税商业、景点接驳、夜间经济；协调器强调游客便利和居民生活之间做分区与时段协调。

用户需求：{need_text}
解析结果：{json.dumps(parse_result, ensure_ascii=False)}
"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "你是多智能体社区规划协同器。只返回可被 json.loads 解析的 JSON 对象。",
            },
            {"role": "user", "content": prompt},
        ],
        "stream": False,
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        result = response.json()
        content = result["choices"][0]["message"]["content"].strip()
        content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"\s*```$", "", content).strip()
        parsed = json.loads(content)
        return normalize_agent_logs(parsed, fallback_logs)
    except Exception as e:
        print(f"Error in DeepSeek Agent API: {str(e)}")
        return fallback_logs


agent_logs = generate_agent_logs(user_input, parse_result)
agent_tabs = st.tabs([agent_name for _, agent_name in AGENT_ORDER])
for tab, (agent_key, _) in zip(agent_tabs, AGENT_ORDER):
    with tab:
        st.markdown(render_agent_content(agent_logs[agent_key]))

st.markdown("## 5. 政策依据与合规解释")

POLICY_FIELD_ALIASES = {
    "policy_basis": ["policy_basis", "政策依据"],
    "compliance_analysis": ["compliance_analysis", "合规性分析"],
    "risk_warning": ["risk_warning", "风险提示"],
    "optimization_suggestions": ["optimization_suggestions", "optimization_suggestion", "优化建议"],
}

POLICY_TITLES = {
    "policy_basis": "政策依据",
    "compliance_analysis": "合规性分析",
    "risk_warning": "风险提示",
    "optimization_suggestions": "优化建议",
}

POLICY_SCENE_KEYWORDS = {
    "候鸟老人友好模式": ["养老", "老人", "医疗", "医院", "社区服务", "15分钟", "慢行", "公园", "菜市场", "无障碍", "公共服务"],
    "年轻家庭模式": ["教育", "学校", "儿童", "家庭", "公共空间", "运动", "社区服务", "交通安全", "居住", "亲子"],
    "游客短租模式": ["旅游", "游客", "短租", "民宿", "酒店", "免税", "景点", "交通接驳", "夜间经济", "公共安全", "三亚", "海南自由贸易港"],
}


def load_policy_documents():
    documents = []
    if not POLICY_DIR.exists():
        return documents

    for file in POLICY_DIR.glob("*.txt"):
        try:
            content = file.read_text(encoding="utf-8", errors="ignore").strip()
            if content:
                documents.append({"source": file.name, "content": content})
        except Exception as e:
            print(f"Error reading policy file {file.name}: {str(e)}")
    return documents


def split_policy_text(content, chunk_size=520):
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", content) if p.strip()]
    if not paragraphs:
        paragraphs = [line.strip() for line in content.splitlines() if line.strip()]

    chunks = []
    for paragraph in paragraphs:
        if len(paragraph) <= chunk_size:
            chunks.append(paragraph)
        else:
            for start in range(0, len(paragraph), chunk_size):
                part = paragraph[start:start + chunk_size].strip()
                if part:
                    chunks.append(part)
    return chunks


def extract_policy_field(block, field_name):
    pattern = rf"【{re.escape(field_name)}】\s*([\s\S]*?)(?=\n【[^】]+】|\Z)"
    match = re.search(pattern, block)
    return clean_policy_text(match.group(1)) if match else ""


def parse_policy_records(content, source_file):
    blocks = [
        block.strip()
        for block in re.split(r"\n\s*\n(?=【政策名称】)", content)
        if "【政策名称】" in block
    ]
    records = []

    for block in blocks:
        policy_name = extract_policy_field(block, "政策名称")
        issuer = extract_policy_field(block, "发布机构")
        year_or_date = extract_policy_field(block, "年份/实施时间") or extract_policy_field(block, "实施时间")
        applicable_scene = extract_policy_field(block, "适用场景")
        snippet = extract_policy_field(block, "政策片段") or clean_policy_text(block)
        support_direction = extract_policy_field(block, "支撑方向")
        if policy_name and snippet:
            records.append(
                {
                    "policy_name": policy_name,
                    "issuer": issuer or "未注明发布机构",
                    "year_or_date": year_or_date or "未注明时间",
                    "applicable_scene": applicable_scene,
                    "snippet": snippet,
                    "support_direction": support_direction or "可作为当前方案合规解释的辅助依据。",
                    "source_file": source_file,
                    "full_text": clean_policy_text(block),
                }
            )

    if records:
        return records

    fallback_records = []
    for chunk in split_policy_text(content):
        fallback_records.append(
            {
                "policy_name": Path(source_file).stem,
                "issuer": "演示政策摘要",
                "year_or_date": "未注明时间",
                "applicable_scene": "",
                "snippet": clean_policy_text(chunk),
                "support_direction": "该摘要片段可作为当前方案合规解释的辅助依据。",
                "source_file": source_file,
                "full_text": clean_policy_text(chunk),
            }
        )
    return fallback_records


def collect_agent_text(agent_logs):
    parts = []
    for agent_key, _ in AGENT_ORDER:
        agent_data = agent_logs.get(agent_key, {})
        parts.extend([str(agent_data.get(field, "")) for field in ["core_need", "reason", "conflict", "compromise"]])
    return " ".join(parts)


def clean_policy_text(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def build_policy_snippet(chunk, matched_keywords, max_length=150):
    cleaned = clean_policy_text(chunk)
    if len(cleaned) <= max_length:
        return cleaned

    hit_positions = [cleaned.find(keyword) for keyword in matched_keywords if keyword and keyword in cleaned]
    first_hit = min(hit_positions) if hit_positions else 0
    start = max(0, first_hit - 42)
    end = min(len(cleaned), start + max_length)
    start = max(0, end - max_length)
    snippet = cleaned[start:end].strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(cleaned):
        snippet = snippet + "..."
    return snippet


def build_policy_relevance_reason(scenario, matched_keywords):
    keyword_text = "、".join(matched_keywords[:5]) if matched_keywords else "当前场景关键词"
    scene_focus = {
        "候鸟老人友好模式": "医疗养老、15分钟生活圈、慢行安全和公共服务可达",
        "年轻家庭模式": "教育可达、儿童友好、公共空间和居住环境",
        "游客短租模式": "旅游服务、交通接驳、短租治理和公共安全",
    }.get(scenario, "社区公共服务、空间治理和合规表达")
    return f"该片段命中「{keyword_text}」等关键词，与{scenario}下的{scene_focus}诉求相关，可支撑本方案的政策依据和优化建议。"


def retrieve_policy_chunks(user_text, parse_result, agent_logs, top_k=5):
    documents = load_policy_documents()
    if not documents:
        return []

    scenario = parse_result.get("scenario", "")
    scenario_keywords = POLICY_SCENE_KEYWORDS.get(scenario, [])
    combined_text = " ".join(
        [
            user_text,
            scenario,
            str(parse_result.get("explanation", "")),
            collect_agent_text(agent_logs),
        ]
    )

    all_keywords = sorted(set(sum(POLICY_SCENE_KEYWORDS.values(), [])))
    matched_keywords = [keyword for keyword in all_keywords if keyword and keyword in combined_text]
    keywords = list(dict.fromkeys(scenario_keywords + matched_keywords))

    scored_chunks = []
    for doc in documents:
        for record in parse_policy_records(doc["content"], doc["source"]):
            searchable = " ".join(
                [
                    record["source_file"],
                    record["policy_name"],
                    record["issuer"],
                    record["year_or_date"],
                    record["applicable_scene"],
                    record["snippet"],
                    record["support_direction"],
                ]
            )
            score = 0
            chunk_matched_keywords = []
            for keyword in keywords:
                count = searchable.count(keyword)
                if count:
                    score += count * (3 if keyword in scenario_keywords else 1)
                    chunk_matched_keywords.append(keyword)
            if score > 0:
                unique_keywords = list(dict.fromkeys(chunk_matched_keywords))
                snippet = build_policy_snippet(record["snippet"], unique_keywords)
                scored_chunks.append(
                    {
                        "policy_name": record["policy_name"],
                        "issuer": record["issuer"],
                        "year_or_date": record["year_or_date"],
                        "source_file": record["source_file"],
                        "source": record["source_file"],
                        "snippet": snippet,
                        "matched_keywords": unique_keywords,
                        "support_direction": record["support_direction"],
                        "relevance_reason": build_policy_relevance_reason(scenario, unique_keywords),
                        "score": score,
                    }
                )

    scored_chunks.sort(key=lambda item: item["score"], reverse=True)
    return scored_chunks[:top_k]


def build_rule_policy_explanation(user_text, parse_result, agent_logs):
    scene_name = parse_result["scenario"]
    need_text = user_text.strip() or "用户暂未输入具体需求，系统根据默认场景进行规划解释。"
    weights = (
        f"医疗{parse_result['medical_weight']:.0%}、教育{parse_result['education_weight']:.0%}、"
        f"商业{parse_result['commerce_weight']:.0%}、绿地{parse_result['green_weight']:.0%}、"
        f"交通{parse_result['traffic_weight']:.0%}"
    )
    coordinator_conflict = agent_logs["coordinator_agent"]["conflict"]

    if scene_name == "候鸟老人友好模式":
        return {
            "sources": [],
            "policy_evidence": [],
            "policy_basis": "对照完整社区和15分钟生活圈建设要求，本方案应重点落实社区养老、基层医疗、菜市场、公园绿地、慢行安全和休息设施配置，并兼顾适老化公共空间与无障碍通行要求。",
            "compliance_analysis": f"用户需求为「{need_text}」，系统识别为候鸟老人友好模式，权重配置为{weights}。医疗、商业、绿地和交通权重共同指向日常服务可达性，符合养老服务、医疗可达和慢行友好导向。",
            "risk_warning": f"当前多Agent冲突显示：{coordinator_conflict}。若商业点位过近或过度引入游客动线，可能影响老人休息、过街安全和慢行连续性。",
            "optimization_suggestions": "建议优先把社区医院、药店、菜场、公园入口和休息座椅布置在老人高频步行路径上；过街点减少绕行并强化安全提示，商业服务采用小尺度、低噪声形态。",
        }
    elif scene_name == "年轻家庭模式":
        return {
            "sources": [],
            "policy_evidence": [],
            "policy_basis": "对照完整社区、儿童友好社区和居住区公共服务设施配置要求，本方案应重点关注教育资源、亲子公共空间、运动活动场地、居住安静和上下学交通安全。",
            "compliance_analysis": f"用户需求为「{need_text}」，系统识别为年轻家庭模式，权重配置为{weights}。教育和生活配套诉求较强，规划解释应优先审查学校可达、儿童活动空间、交通组织和公共空间均衡。",
            "risk_warning": f"当前多Agent冲突显示：{coordinator_conflict}。若商业、培训和亲子消费过度集中，可能造成接送拥堵、噪声外溢和住宅安静度下降。",
            "optimization_suggestions": "建议以学校和亲子空间为核心组织安全步行线路，商业和培训/文体服务布置在外围或交通节点；通过分时管理、慢行优先和公共空间共享提升家庭生活便利。",
        }
    else:
        return {
            "sources": [],
            "policy_evidence": [],
            "policy_basis": "对照海南自贸港、国际旅游消费中心、旅游公共服务和社区治理相关要求，本方案应重点关注旅游消费、交通接驳、短租治理、夜间经济、公共安全和居民生活秩序。",
            "compliance_analysis": f"用户需求为「{need_text}」，系统识别为游客短租模式，权重配置为{weights}。商业和交通权重较高，说明规划应优先审查民宿酒店、免税商业、景点接驳和夜间出行的承载关系。",
            "risk_warning": f"当前多Agent冲突显示：{coordinator_conflict}。若游客动线、夜间经济和短租客流缺少边界，可能挤占居民公共空间、带来噪声扰动和治理压力。",
            "optimization_suggestions": "建议采用游客服务分区、公交接驳优化、夜间时段管理和公共空间承载控制；将免税商业、民宿酒店和景点换乘节点集中到适合承载的界面，保护居民日常生活秩序。",
        }


def pick_policy_field(policy_value, field_name, fallback_value):
    if isinstance(policy_value, dict):
        for alias in POLICY_FIELD_ALIASES[field_name]:
            value = policy_value.get(alias)
            if value is not None and str(value).strip():
                return str(value).strip()
    return fallback_value


def normalize_policy_explanation(raw_result, fallback_policy):
    if not isinstance(raw_result, dict):
        return fallback_policy
    sources = raw_result.get("sources", fallback_policy.get("sources", []))
    if isinstance(sources, str):
        sources = [sources]
    if not isinstance(sources, list):
        sources = fallback_policy.get("sources", [])

    fallback_sources = fallback_policy.get("sources", [])
    if fallback_sources:
        allowed_sources = set(fallback_sources)
        sources = [source for source in sources if source in allowed_sources] or fallback_sources

    normalized = {
        "sources": [str(source).strip() for source in sources if str(source).strip()],
        "policy_evidence": fallback_policy.get("policy_evidence", []),
    }
    normalized.update({
        field: pick_policy_field(raw_result, field, fallback_policy[field])
        for field in POLICY_FIELD_ALIASES
    })
    return normalized


def get_policy_evidence(policy_sections):
    if not isinstance(policy_sections, dict):
        return []
    evidence = policy_sections.get("policy_evidence", [])
    return evidence if isinstance(evidence, list) else []


def render_policy_evidence_markdown(policy_sections):
    evidence_items = get_policy_evidence(policy_sections)
    if not evidence_items:
        return "## 政策证据链\n\n当前未检索到明确政策片段，系统使用规则模板生成合规解释。"

    blocks = ["## 政策证据链"]
    for index, item in enumerate(evidence_items[:5], start=1):
        keywords = "、".join(item.get("matched_keywords", [])[:8]) or "无明确关键词"
        policy_name = item.get("policy_name") or item.get("source") or "未知政策"
        issuer = item.get("issuer", "未注明发布机构")
        year_or_date = item.get("year_or_date", "未注明时间")
        source_file = item.get("source_file") or item.get("source", "未知来源文件")
        blocks.append(
            f"""#### 政策{index}：{policy_name}

- 发布机构：{issuer}
- 时间：{year_or_date}
- 来源文件：{source_file}
- 命中关键词：{keywords}
- 命中片段：{item.get('snippet', '未检索到明确片段')}
- 支撑方向：{item.get('support_direction', '该政策片段可作为当前方案合规解释的辅助依据。')}"""
        )
    return "\n\n".join(blocks)


def render_policy_evidence_cards(policy_sections):
    evidence_items = get_policy_evidence(policy_sections)
    st.markdown("### 政策证据链")
    if not evidence_items:
        st.markdown(
            """
<div class="policy-evidence-empty">
当前未检索到明确政策片段，系统使用规则模板生成合规解释。
</div>
""",
            unsafe_allow_html=True,
        )
        return

    cards = []
    for item in evidence_items[:5]:
        policy_name = html.escape(str(item.get("policy_name") or item.get("source") or "未知政策"))
        issuer = html.escape(str(item.get("issuer", "未注明发布机构")))
        year_or_date = html.escape(str(item.get("year_or_date", "未注明时间")))
        source_file = html.escape(str(item.get("source_file") or item.get("source", "未知来源文件")))
        snippet = html.escape(str(item.get("snippet", "未检索到明确片段")))
        support_direction = html.escape(str(item.get("support_direction", "该政策片段可作为当前方案合规解释的辅助依据。")))
        keywords = item.get("matched_keywords", [])
        if not isinstance(keywords, list):
            keywords = []
        keyword_badges = "".join(
            [f'<span class="policy-keyword-badge">{html.escape(str(keyword))}</span>' for keyword in keywords[:8]]
        ) or '<span class="policy-keyword-badge">场景相关</span>'
        cards.append(
            f"""
<div class="policy-evidence-card">
  <div class="policy-evidence-source">{policy_name}</div>
  <p class="policy-evidence-meta">发布机构 / 年份：{issuer} / {year_or_date}</p>
  <p class="policy-evidence-meta">来源文件：{source_file}</p>
  <div class="policy-evidence-label">命中关键词</div>
  <div class="policy-keyword-badges">{keyword_badges}</div>
  <div class="policy-evidence-label">命中片段</div>
  <p class="policy-evidence-snippet">{snippet}</p>
  <div class="policy-evidence-label">支撑本方案的方向</div>
  <p class="policy-evidence-reason">{support_direction}</p>
</div>
"""
        )

    st.markdown(
        f"""
<div class="policy-evidence-grid">
{''.join(cards)}
</div>
""",
        unsafe_allow_html=True,
    )


def render_policy_markdown(policy_sections):
    source_lines = "\n".join([f"- {source}" for source in policy_sections.get("sources", [])])
    source_text = f"### 已检索政策来源\n\n{source_lines or '未检索到本地政策片段，使用规则模板解释。'}"
    evidence_text = render_policy_evidence_markdown(policy_sections)
    section_text = "\n\n".join(
        [f"### {POLICY_TITLES[field]}\n\n{policy_sections[field]}" for field in POLICY_FIELD_ALIASES]
    )
    return f"{source_text}\n\n{evidence_text}\n\n{section_text}"


def render_policy_card(title, content):
    safe_title = html.escape(title)
    content_text = str(content).strip() if content is not None else ""
    if not content_text or content_text.lower() in ["none", "null", "nan"]:
        content_text = "系统已根据当前场景生成默认政策解释。"
    safe_content = html.escape(content_text)
    return f"""
<div class="policy-card">
  <h3>{safe_title}</h3>
  <p>{safe_content}</p>
</div>
"""


def render_policy_card_grid(policy_sections):
    cards_html = "".join(
        [
            render_policy_card(POLICY_TITLES[field], policy_sections.get(field, ""))
            for field in POLICY_FIELD_ALIASES
        ]
    )
    st.markdown(
        f"""
<div class="policy-card-grid">
{cards_html}
</div>
""",
        unsafe_allow_html=True,
    )


def generate_policy_explanation(user_text, parse_result, agent_logs):
    fallback_policy = build_rule_policy_explanation(user_text, parse_result, agent_logs)
    retrieved_chunks = retrieve_policy_chunks(user_text, parse_result, agent_logs, top_k=5)
    retrieved_sources = sorted({item["source"] for item in retrieved_chunks})
    fallback_policy["sources"] = retrieved_sources
    fallback_policy["policy_evidence"] = retrieved_chunks
    api_key = get_config_value("DEEPSEEK_API_KEY")
    base_url = get_config_value("DEEPSEEK_BASE_URL", "https://api.sydney-ai.com/v1")
    model = get_config_value("DEEPSEEK_MODEL", "deepseek-chat")
    if not api_key or not retrieved_chunks:
        return fallback_policy

    need_text = user_text.strip() or "用户暂未输入具体需求，系统根据默认场景进行规划解释。"
    agent_conflicts = {
        REPORT_AGENT_NAMES.get(agent_key, agent_name): agent_logs[agent_key]["conflict"]
        for agent_key, agent_name in AGENT_ORDER
    }
    weights = {field: parse_result[field] for field in WEIGHT_FIELDS}
    policy_context = "\n\n".join(
        [
            f"政策名称：{item.get('policy_name', '未知政策')}｜发布机构：{item.get('issuer', '未注明发布机构')}｜时间：{item.get('year_or_date', '未注明时间')}｜来源文件：{item.get('source_file', item.get('source', '未知来源'))}｜相关度：{item['score']}｜命中关键词：{'、'.join(item.get('matched_keywords', []))}\n{item['snippet']}\n支撑方向：{item.get('support_direction', item.get('relevance_reason', '该政策片段可作为当前方案合规解释的辅助依据。'))}"
            for item in retrieved_chunks
        ]
    )
    url = base_url.rstrip("/") + "/chat/completions"
    prompt = f"""
请基于用户需求、AI识别场景、指标权重、多Agent冲突和检索到的政策片段，生成更像真实规划审查的政策解释 JSON。
只返回 JSON，不要 Markdown，不要解释文字，不要代码块。

JSON 必须严格包含以下字段，且每个字段不能为空：
{{
  "sources": ["政策文件1.txt", "政策文件2.txt"],
  "policy_basis": "政策依据",
  "compliance_analysis": "合规性分析",
  "risk_warning": "风险提示",
  "optimization_suggestions": "优化建议"
}}

场景侧重点：
1. 候鸟老人友好模式：重点解释15分钟生活圈、社区养老、医疗可达、慢行安全、休息设施。
2. 年轻家庭模式：重点解释教育资源、儿童友好、公共空间、居住安静、交通安全。
3. 游客短租模式：重点解释海南自贸港、旅游消费、交通接驳、短租治理、夜间经济、居民生活秩序。

用户原始需求：{need_text}
AI识别结果：{json.dumps(parse_result, ensure_ascii=False)}
指标权重：{json.dumps(weights, ensure_ascii=False)}
多Agent冲突：{json.dumps(agent_conflicts, ensure_ascii=False)}
检索到的政策片段：
{policy_context}
"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "你是社区规划合规审查助手。只返回可被 json.loads 解析的 JSON 对象。",
            },
            {"role": "user", "content": prompt},
        ],
        "stream": False,
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        result = response.json()
        content = result["choices"][0]["message"]["content"].strip()
        content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"\s*```$", "", content).strip()
        parsed = json.loads(content)
        return normalize_policy_explanation(parsed, fallback_policy)
    except Exception as e:
        print(f"Error in DeepSeek Policy API: {str(e)}")
        return fallback_policy


policy_sections = generate_policy_explanation(user_input, parse_result, agent_logs)
policy_explanation = render_policy_markdown(policy_sections)

policy_evidence_items = get_policy_evidence(policy_sections)
if policy_evidence_items:
    source_count = len(policy_evidence_items)
    source_items = "".join(
        [
            f"<li>{html.escape(str(item.get('policy_name', '未知政策')))}｜{html.escape(str(item.get('issuer', '未注明发布机构')))}｜{html.escape(str(item.get('year_or_date', '未注明时间')))}</li>"
            for item in policy_evidence_items[:5]
        ]
    )
    st.markdown(
        f"""
<div class="policy-source-card">
  <strong>已检索到 {source_count} 条结构化政策证据</strong>
  <ul>{source_items}</ul>
</div>
""",
        unsafe_allow_html=True,
    )
else:
    st.markdown(
        """
<div class="policy-source-card">
  <strong>未检索到明确政策来源，已使用内置规则模板生成解释</strong>
</div>
""",
        unsafe_allow_html=True,
    )

render_policy_evidence_cards(policy_sections)
render_policy_card_grid(policy_sections)

st.markdown("## 6. 一键生成规划报告")

def safe_report_text(value, fallback="暂无说明"):
    if value is None:
        return fallback
    text = str(value).strip()
    if not text or text.lower() in ["none", "null", "nan"]:
        return fallback
    return text


def get_weight_explanation(scene_name):
    if scene_name == "候鸟老人友好模式":
        return "本方案更关注医疗、绿地和慢行交通可达性，适合候鸟老人日常就医、买菜、休闲和安全步行需求。"
    if scene_name == "年轻家庭模式":
        return "本方案更关注教育、公共空间和居住安静，适合年轻家庭的通学、亲子活动和日常生活配套需求。"
    return "本方案更关注交通接驳、商业消费和景点服务便利性，适合游客短租、免税消费和夜间出行场景。"


def summarize_space_statistics(df):
    if df is None:
        return "当前演示版本读取已生成规划结果，空间统计摘要暂不可用。"

    try:
        row_count = len(df)
        col_count = len(df.columns)
        numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
        area_cols = [col for col in df.columns if any(key in str(col).lower() for key in ["area", "面积", "用地", "绿地"])]
        length_cols = [col for col in df.columns if any(key in str(col).lower() for key in ["length", "长度", "道路", "road"])]

        parts = [f"- 统计记录数量：{row_count} 条。", f"- 指标字段数量：{col_count} 个。"]
        if area_cols:
            area_text = "、".join([str(col) for col in area_cols[:4]])
            parts.append(f"- 主要面积相关字段：{area_text}。")
        if length_cols:
            length_text = "、".join([str(col) for col in length_cols[:4]])
            parts.append(f"- 主要长度/道路相关字段：{length_text}。")
        if numeric_cols:
            numeric_text = "、".join([str(col) for col in numeric_cols[:5]])
            parts.append(f"- 可用于数值图展示的字段：{numeric_text}。")
        return "\n".join(parts)
    except Exception:
        return "当前演示版本读取已生成规划结果，空间统计摘要暂不可用。"


def get_report_agent_logs(agent_logs, user_text, parse_result):
    fallback_agents = build_rule_agent_logs(user_text, parse_result)
    report_agents = {}
    for agent_key, _ in AGENT_ORDER:
        current = agent_logs.get(agent_key, {}) if isinstance(agent_logs, dict) else {}
        fallback_agent = fallback_agents[agent_key]
        report_agents[agent_key] = {
            "core_need": safe_report_text(current.get("core_need"), fallback_agent["core_need"]),
            "reason": safe_report_text(current.get("reason"), fallback_agent["reason"]),
            "conflict": safe_report_text(current.get("conflict"), fallback_agent["conflict"]),
            "compromise": safe_report_text(current.get("compromise"), fallback_agent["compromise"]),
        }
    return report_agents


def render_agent_report(agent_logs):
    blocks = []
    for agent_key, agent_name in AGENT_ORDER:
        report_name = REPORT_AGENT_NAMES.get(agent_key, agent_name)
        agent_data = agent_logs[agent_key]
        blocks.append(
            f"""### {report_name}

- 核心诉求：{agent_data['core_need']}
- 支持理由：{agent_data['reason']}
- 潜在冲突：{agent_data['conflict']}
- 妥协建议：{agent_data['compromise']}"""
        )
    return "\n\n".join(blocks)


def render_policy_report(policy_sections, user_text, parse_result, agent_logs):
    fallback_policy = build_rule_policy_explanation(user_text, parse_result, agent_logs)
    evidence_items = get_policy_evidence(policy_sections if isinstance(policy_sections, dict) else {})
    sources = policy_sections.get("sources", []) if isinstance(policy_sections, dict) else []
    if evidence_items:
        source_text = "\n".join(
            [
                f"- {safe_report_text(item.get('policy_name'), '未知政策')}｜{safe_report_text(item.get('issuer'), '未注明发布机构')}｜{safe_report_text(item.get('year_or_date'), '未注明时间')}"
                for item in evidence_items[:5]
            ]
        )
    elif sources:
        source_text = "\n".join([f"- {safe_report_text(source, '未知政策来源')}" for source in sources])
    else:
        source_text = "未检索到明确政策来源，系统使用内置规则模板生成政策解释。"

    evidence_text = render_policy_evidence_markdown(policy_sections if isinstance(policy_sections, dict) else {})
    policy_basis = safe_report_text(policy_sections.get("policy_basis") if isinstance(policy_sections, dict) else None, fallback_policy["policy_basis"])
    compliance = safe_report_text(policy_sections.get("compliance_analysis") if isinstance(policy_sections, dict) else None, fallback_policy["compliance_analysis"])
    risk = safe_report_text(policy_sections.get("risk_warning") if isinstance(policy_sections, dict) else None, fallback_policy["risk_warning"])
    suggestion = safe_report_text(policy_sections.get("optimization_suggestions") if isinstance(policy_sections, dict) else None, fallback_policy["optimization_suggestions"])

    return f"""### 已检索政策来源

{source_text}

{evidence_text}

### 政策依据

{policy_basis}

### 合规性分析

{compliance}

### 风险提示

{risk}

### 优化建议

{suggestion}"""


def get_integrated_risks_and_suggestions(scene_name):
    if scene_name == "候鸟老人友好模式":
        return (
            ["慢行安全不足，老人日常步行和过街存在安全隐患。", "医疗与养老服务可达性不足，可能影响候鸟老人日常照护效率。", "商业设施过度靠近居住与休憩空间，可能干扰安静生活。"],
            ["增设休息座椅、无障碍路径和安全过街节点。", "强化社区卫生服务、药店、养老服务与日常生活服务联动。", "将高活力商业活动控制在合适范围，保留安静、连续的慢行生活空间。"],
        )
    if scene_name == "年轻家庭模式":
        return (
            ["教育资源不足，可能影响年轻家庭的就近入学便利性。", "上下学时段交通安全风险较高，接送动线可能造成拥堵。", "商业和活动空间过度集中，可能影响居住安静。"],
            ["强化学校、托育、亲子空间和运动活动场地配置。", "优化人车分流、安全步行路线和接送交通组织。", "平衡运动空间、商业空间和居住空间，控制噪声外溢。"],
        )
    return (
        ["夜间噪声和游客活动可能引发居民投诉。", "短租民宿和流动人口带来社区治理压力。", "景点、商业区和接驳节点可能出现交通拥堵。"],
        ["设置游客活动集中区，减少游客动线对居民生活区的穿越。", "加强夜间管理、安静时段规则和短租经营秩序引导。", "优化公交接驳、慢行路线和停车组织，提升旅游出行效率。"],
    )


def build_report():
    report_scene = safe_report_text(parse_result.get("scenario"), final_scene)
    report_method = safe_report_text(parse_result.get("parse_method"), "规则解析")
    report_walking_time = safe_report_text(parse_result.get("walking_time"), "15")
    report_explanation = safe_report_text(parse_result.get("explanation"), "系统根据当前场景和用户需求完成需求解析。")
    report_user_need = safe_report_text(user_input, "用户未输入明确需求，系统使用当前选择场景进行默认分析。")

    weights_text = "\n".join([f"- {k}: {v:.0%}" for k, v in parsed_weight_items])
    weight_total = sum([value for _, value in parsed_weight_items])
    weight_explanation = get_weight_explanation(report_scene)
    planning_summary = summarize_space_statistics(summary_df)
    report_agent_logs = get_report_agent_logs(agent_logs, user_input, parse_result)
    logs_text = render_agent_report(report_agent_logs)
    policy_text = render_policy_report(policy_sections, user_input, parse_result, report_agent_logs)
    instant_plan_text = render_instant_plan_report(report_scene, parse_result)
    risks, suggestions = get_integrated_risks_and_suggestions(report_scene)
    risks_text = "\n".join([f"{index}. {item}" for index, item in enumerate(risks, start=1)])
    suggestions_text = "\n".join([f"{index}. {item}" for index, item in enumerate(suggestions, start=1)])
    image_note = f"当前场景对应的规划结果图为 `{img_path.name}`。报告不嵌入图片，页面端展示上游 PPO/SGNN 项目离线生成结果的示意图和统计表，不进行现场训练，也不实时推理生成新规划图。当前页面图片为展示端基于 GeoJSON `type` 字段生成的可视化结果，颜色采用展示端渲染配色。页面端支持基于离线 GeoJSON 的动态渲染与场景重点类型高亮，但当前版本不进行现场 PPO/SGNN 训练或实时生成新规划结果。"
    legend_report_rows = [
        "| type | 上游类型 | 中文含义 | 展示端颜色 | 说明 |",
        "|---|---|---|---|---|",
    ]
    legend_report_rows.extend(
        [
            f"| {item['编号']} | {item['上游类型']} | {item['类型名称']} | {item['展示端颜色']} | {item['说明']} |"
            for item in LAND_USE_LEGEND
        ]
    )
    legend_report_text = "\n".join(legend_report_rows)
    
    report = f"""# 琼岛智划规划分析报告

生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

## 一、项目背景

本系统面向海南自贸港与三亚滨海社区规划场景，采用“离线规划引擎 + 在线智能解释”架构：展示端读取上游 PPO/SGNN 项目离线生成的规划结果图和统计表，不进行现场训练；DeepSeek、多Agent、RAG 模块用于需求解析、解释生成和合规辅助，不直接改变底层空间规划结果。

## 二、用户需求

{report_user_need}

## 三、AI识别结果

- 识别场景：{report_scene}
- 解析方式：{report_method}
- 建议步行时间：{report_walking_time} 分钟
- 解析说明：{report_explanation}

## 四、指标权重分析

{weights_text}

五项权重合计约为 {weight_total:.0%}。{weight_explanation}

## 五、规划结果说明

- 核心规划目标：{safe_report_text(config.get("main_goal"), "围绕当前识别场景优化社区空间服务能力。")}
- 推荐方案：{safe_report_text(config.get("recommendation"), "优先展示当前场景对应的已生成规划结果，并结合需求解析进行说明。")}
- 规划图说明：{image_note}
- 空间统计结果摘要：

{planning_summary}

### 图例说明

type 含义来自上游 DRL urban planning 项目的 `city_config.py`。当前页面图片为展示端基于 GeoJSON `type` 字段生成的可视化结果，颜色采用展示端渲染配色。

{legend_report_text}

注：颜色为展示端渲染配色，土地类型含义以上游 `city_config.py` 的 `type` 定义为准。其中 `type=4` 表示**居住用地**，`type=7`、`type=8` 表示**绿地类用地**。

{instant_plan_text}

## 六、离线规划引擎 + 在线智能解释（PPO-GNN/SGNN离线规划流程说明）

上游 PPO-GNN/SGNN 项目作为底层算法依据，以社区空间地块、道路、设施和用地类型等空间属性为输入，将地块或空间单元抽象为图节点，并把邻接关系、道路连接或空间相邻关系建模为边。当前琼岛智划展示端没有集成完整 PPO/SGNN 训练引擎，也不在页面端执行模型训练或实时推理生成新规划图。

当前演示端读取上游 PPO/SGNN 项目离线生成的规划结果图和统计表，不进行现场训练；展示端结合用户需求解析、多Agent解释辅助和 RAG 合规辅助，生成可展示、可下载的规划分析结果。DeepSeek、多Agent、RAG 模块用于需求解析、解释生成和合规辅助，不直接改变底层空间规划结果。

## 七、多Agent解释辅助

{logs_text}

## 八、RAG合规辅助

{policy_text}

## 九、综合风险与优化建议

### 主要风险

{risks_text}

### 优化建议

{suggestions_text}

## 十、结论

本次规划建议以 DeepSeek/LLM 对自然语言需求的解析为入口，结合多Agent对居民、政府、商业三方诉求的解释辅助，并通过轻量级 RAG 对本地政策文本进行合规解释支撑。整体方案面向海南自贸港社区治理与三亚滨海社区更新场景，能够为15分钟生活圈目标解释、公共服务配置说明和空间治理展示提供可解释依据；这些在线智能解释模块不直接改变底层空间规划结果。

## 十一、当前演示说明

当前版本为第一版演示系统，展示端读取上游 PPO/SGNN 项目离线生成的空间规划结果图、统计表和本地政策文本，不进行现场模型训练，也不实时推理生成新规划图。上游算法可在 Linux / WSL2 环境独立复现或扩展，展示端运行于 Windows + Streamlit 环境。

---

琼岛智划 v0.1 | 面向海南自贸港的 AI 多智能体社区规划系统
"""
    return report


def clean_markdown_inline(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    return text.strip()


def parse_markdown_table_row(line):
    cells = [clean_markdown_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
    return cells


def is_markdown_table_separator(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_markdown_table(document, table_lines):
    rows = [
        parse_markdown_table_row(line)
        for line in table_lines
        if not is_markdown_table_separator(line)
    ]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_values in rows:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = row_values[index] if index < len(row_values) else ""


def generate_report_docx(report_markdown: str) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    document = Document()
    document.core_properties.title = "琼岛智划规划分析报告"
    document.core_properties.author = "琼岛智划"

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            document.styles[style_name].font.name = "Microsoft YaHei"
        except KeyError:
            pass

    lines = report_markdown.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while index < len(lines):
                candidate = lines[index].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_lines.append(candidate)
                    index += 1
                else:
                    break
            add_markdown_table(document, table_lines)
            continue

        if line.startswith("# "):
            document.add_heading(clean_markdown_inline(line[2:]), level=0)
        elif line.startswith("## "):
            document.add_heading(clean_markdown_inline(line[3:]), level=1)
        elif line.startswith("### "):
            document.add_heading(clean_markdown_inline(line[4:]), level=2)
        elif line.startswith("#### "):
            document.add_heading(clean_markdown_inline(line[5:]), level=3)
        elif line.startswith("- "):
            document.add_paragraph(clean_markdown_inline(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(clean_markdown_inline(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        elif line == "---":
            document.add_paragraph("")
        else:
            document.add_paragraph(clean_markdown_inline(line))

        index += 1

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


report = build_report()

st.download_button(
    label="下载规划分析报告 Markdown",
    data=report,
    file_name="琼岛智划_规划分析报告.md",
    mime="text/markdown"
)

if DOCX_AVAILABLE:
    try:
        report_docx = generate_report_docx(report)
        st.download_button(
            label="下载规划分析报告 Word",
            data=report_docx,
            file_name="琼岛智划规划分析报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.warning(f"Word 报告生成失败：{e}")
else:
    st.warning("Word 导出依赖 python-docx 未安装，请先安装依赖。")

with st.expander("预览报告内容"):
    st.markdown(report)

st.divider()

st.caption(
    "琼岛智划 v0.1 | 离线规划引擎 + 在线智能解释 | 当前展示端不进行现场训练。"
)
